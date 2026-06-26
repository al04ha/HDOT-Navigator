"""
doc_extractor.py
Extracts plain text and detailed formatting reports from DOCX, PDF, and XLSX files.
"""

from docx import Document
from docx.oxml.ns import qn
from pypdf import PdfReader
import openpyxl


def fmt_bool(val) -> str:
    """Convert a tri-state bool (True / False / None) to a readable string."""
    return "inherited" if val is None else str(val)


def extract_formatting_report(file_path: str) -> str:
    """
    Walk every paragraph and run in a DOCX file and return a structured
    text representation of all formatting attributes.
    Used as input to the AI formatting checker.
    """
    doc = Document(file_path)
    lines = []

    for para in doc.paragraphs:
        if not para.text.strip():
            lines.append("[BLANK LINE]")
            continue

        fmt        = para.paragraph_format
        style_name = para.style.name if para.style else "Normal"
        align      = str(para.alignment) if para.alignment else "inherited"

        try:
            space_before = f"{fmt.space_before.pt}pt" if fmt.space_before else "inherited"
        except Exception:
            space_before = "inherited"
        try:
            space_after = f"{fmt.space_after.pt}pt" if fmt.space_after else "inherited"
        except Exception:
            space_after = "inherited"
        try:
            line_spacing = str(fmt.line_spacing) if fmt.line_spacing else "inherited"
        except Exception:
            line_spacing = "inherited"

        lines.append(
            f"[PARA style='{style_name}' align={align} "
            f"space_before={space_before} space_after={space_after} "
            f"line_spacing={line_spacing}]"
        )

        for run in para.runs:
            if not run.text:
                continue
            run_info = f"  [RUN text={repr(run.text)}"
            run_info += f" bold={fmt_bool(run.bold)}"
            run_info += f" italic={fmt_bool(run.italic)}"
            run_info += f" underline={fmt_bool(run.underline)}"
            try:
                if run.font.size:
                    run_info += f" size={run.font.size.pt}pt"
            except Exception:
                pass
            try:
                if run.font.color and run.font.color.type is not None:
                    run_info += f" color={run.font.color.rgb}"
            except Exception:
                pass
            try:
                if run.font.name:
                    run_info += f" font='{run.font.name}'"
            except Exception:
                pass
            run_info += "]"
            lines.append(run_info)

    for ti, table in enumerate(doc.tables):
        lines.append(f"[TABLE {ti + 1}]")
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if not para.text.strip():
                        continue
                    lines.append(f"  [CELL text={repr(para.text.strip())}]")
                    for run in para.runs:
                        if not run.text.strip():
                            continue
                        run_info = f"    [RUN text={repr(run.text)}"
                        run_info += f" bold={fmt_bool(run.bold)}"
                        run_info += f" italic={fmt_bool(run.italic)}"
                        run_info += f" underline={fmt_bool(run.underline)}"
                        try:
                            if run.font.size:
                                run_info += f" size={run.font.size.pt}pt"
                        except Exception:
                            pass
                        try:
                            if run.font.color and run.font.color.type is not None:
                                run_info += f" color={run.font.color.rgb}"
                        except Exception:
                            pass
                        try:
                            if run.font.name:
                                run_info += f" font='{run.font.name}'"
                        except Exception:
                            pass
                        run_info += "]"
                        lines.append(run_info)

    return "\n".join(lines)


def extract_text_from_file(file_path: str) -> str:
    """
    Extract plain text content from a DOCX, PDF, or XLSX file.
    Used for indexing into ChromaDB.
    """
    import os
    _, ext = os.path.splitext(file_path.lower())
    try:
        if ext == ".pdf":
            reader = PdfReader(file_path)
            return "\n".join(p.extract_text() for p in reader.pages if p.extract_text())

        if ext == ".docx":
            doc = Document(file_path)
            parts = []

            def para_text(elem):
                run_texts = []
                for r in elem.findall(f'.//{qn("w:r")}'):
                    t_nodes = r.findall(qn("w:t"))
                    run_str = "".join(t.text for t in t_nodes if t.text)
                    if not run_str:
                        continue
                    if run_str.strip().isdigit() and len(run_str.strip()) <= 6:
                        continue
                    run_texts.append(run_str)
                if not run_texts:
                    return ""
                result = run_texts[0]
                for rt in run_texts[1:]:
                    result += rt if (result.endswith(" ") or rt.startswith(" ")) else " " + rt
                return result.strip()

            def table_to_kv(tbl_elem):
                kv_lines = []
                for row in tbl_elem.findall(f'.//{qn("w:tr")}'):
                    cells = row.findall(f'.//{qn("w:tc")}')
                    if not cells:
                        continue
                    cell_paras = []
                    for cell in cells:
                        paras = [para_text(p) for p in cell.findall(f'.//{qn("w:p")}') if para_text(p)]
                        cell_paras.append(paras)
                    if len(cell_paras) == 1:
                        kv_lines.extend(cell_paras[0])
                        continue
                    label_paras = cell_paras[0]
                    value_paras = cell_paras[1]
                    extra_paras = cell_paras[2] if len(cell_paras) > 2 else []
                    for i, lbl_para in enumerate(label_paras):
                        sub_labels = [s.strip().rstrip(":") for s in lbl_para.split(":") if s.strip()]
                        for j, sub_lbl in enumerate(sub_labels):
                            val = value_paras[-1] if (j == len(sub_labels) - 1 and len(sub_labels) > 1) else (value_paras[i] if i < len(value_paras) else "")
                            kv_lines.append(f"{sub_lbl}: {val}")
                    for ep in extra_paras:
                        kv_lines.append(ep)
                return kv_lines

            def process_block(elem):
                tag = elem.tag.split("}")[-1] if "}" in elem.tag else elem.tag
                if tag == "p":
                    t = para_text(elem)
                    if t:
                        parts.append(t)
                elif tag == "tbl":
                    parts.extend(table_to_kv(elem))
                elif tag == "sdt":
                    content = elem.find(qn("w:sdtContent"))
                    if content is not None:
                        for child in content:
                            process_block(child)
                elif tag in ("txbxContent", "drawing", "pict"):
                    t = "".join(x.text for x in elem.iter(qn("w:t")) if x.text).strip()
                    if t:
                        parts.append(t)
                else:
                    for child in elem:
                        ct = child.tag.split("}")[-1] if "}" in child.tag else child.tag
                        if ct in ("p", "tbl", "sdt", "txbxContent", "drawing", "pict"):
                            process_block(child)

            for block in doc.element.body:
                process_block(block)
            for section in doc.sections:
                for hf in [section.header, section.footer,
                           section.first_page_header, section.first_page_footer]:
                    try:
                        if hf and not hf.is_linked_to_previous:
                            for child in hf._element:
                                process_block(child)
                    except Exception:
                        pass
            return "\n".join(parts)

        if ext in (".xlsx", ".xlsm"):
            wb = openpyxl.load_workbook(file_path, data_only=True, read_only=True)
            rows = []
            for sheet in wb.sheetnames:
                rows.append(f"--- Sheet: {sheet} ---")
                for row in wb[sheet].iter_rows(values_only=True):
                    r = " | ".join(str(c) for c in row if c is not None)
                    if r.strip():
                        rows.append(r)
            return "\n".join(rows)

    except Exception as e:
        print(f"[extract_text] Error reading {file_path}: {e}")
    return ""


def chunk_text(text: str, chunk_size: int = 1000, overlap: int = 200) -> list:
    """Split text into overlapping word chunks for embedding."""
    if not text.strip():
        return []
    words = text.split()
    chunks = []
    for i in range(0, len(words), chunk_size - overlap):
        chunks.append(" ".join(words[i:i + chunk_size]))
        if i + chunk_size >= len(words):
            break
    return chunks
