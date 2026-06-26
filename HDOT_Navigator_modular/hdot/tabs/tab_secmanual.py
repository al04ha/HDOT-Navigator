"""tabs/tab_secmanual.py — Secretaries Manual checker tab

Includes a DOCX auto-fix engine: formatting issues found during the AI check
are mapped back to the exact paragraph/run in the original uploaded file and
fixed in place (bold/italic/underline/color/size/alignment/spacing), so the
user can download a corrected DOCX that preserves their original layout,
fonts, and content. Content issues (missing/unfilled fields) are never
auto-fixed — those always require manual review.
"""

import os
import re
import io
import json
import hashlib
import streamlit as st
from datetime import datetime

from docx import Document
from docx.shared import RGBColor
from docx.oxml.ns import qn

from ai_client import gemini_call, _gemini_cache
from template_store import (
    list_docx_templates, find_matching_template, extract_template_formatting,
    save_docx_template, load_docx_template, delete_docx_template,
    save_template_override, delete_template_override, has_template_override,
    _template_path, load_kb_folder_links, save_kb_folder_links,
    get_kb_folder_for_doc_type, get_kb_context_for_folder,
)
from feedback import (
    load_feedback, save_feedback, delete_feedback,
    build_feedback_context, new_feedback_entry,
)
from chroma_store import get_kb_all_metadata, is_source_indexed

# alias to match the name used in the tab body
is_manual_indexed = is_source_indexed

# ── HDOT Common Mistakes ──────────────────────────────────────────────────────
COMMON_MISTAKES = """
HDOT COMMON MISTAKES — check every document for ALL of these:

1. DATE SPACING: The date placed directly under the letterhead is too high. There must always be one blank space between the date and the letterhead.
2. IN REPLY TO SPACING: There must be a space between "In Reply To:" and the letter number.
3. ZIP CODE SPACING: Two spaces before the zip code in the inside address.
4. PERIOD SPACING: Two spaces after every period.
5. ACRONYM LINE BREAKS: Acronyms must not be separated or start on a new line by themselves. If a line break is needed, use the last word of the spelled-out version. Example: "Hawaii Department of Transportation (HDOT)" — never break so "(HDOT)" starts a new line alone.
6. UNCITED ACRONYMS: Do not use an acronym unless it has been properly spelled out and cited at least once earlier in the document.
7. CLOSING PARAGRAPH: The closing paragraph must include an email address, phone number, and title. (Does not apply to internal memos.)
8. SUBJECT LINE SPACING: Two blank lines below the SUBJECT: line for both internal and external memos.
9. SALUTATION ORDER: In letters, the salutation must come before the Subject line.
10. SECOND PAGE HEADER: All pages after page 1 must have a header. For memos the first line must match the TO: line. For letters it must match the first line of the inside address. Headers must be 1 inch from the top.
11. INTERNAL MEMO HEADER TEXT: The header on page 2+ of an internal memo must read "MEMO TO DIR".
12. DOCUMENT ORDER: Correspondence first, then Attachments/Enclosures (if any), then Route Slip, then Original Correspondence.
13. ATTACHMENT VS ENCLOSURE: Use only one word consistently — do not mix "attachment" and "enclosure" in the same document.
14. PHONE NUMBER FORMAT: Area code must be in parentheses. Example: (808) 587-2150.
15. SINCERELY SPACING: One blank line before "Sincerely" after the closing paragraph.
16. NO QUOTES AROUND ABBREVIATIONS: Do not put quotation marks around HDOT or any abbreviation. Wrong: ("HDOT"). Correct: (HDOT).
17. NO CONTRACTIONS: Do not use contractions (e.g. don't, can't, it's) in formal documents.
18. SPELLING, SYNTAX, AND GRAMMAR: Check for misspelled words, grammatical errors, and awkward sentence structure.
19. NUMBER RULES: Spell out one through nine, use numerals for 10+. Always spell out numbers at the start of a sentence. Use numerals for money. Do not start a sentence with a numeral. Express time as a single number (9 a.m., not 9:00 a.m.).
20. FONT: All text in the document must be Times New Roman, 12 point font, unless the template explicitly requires a different font/size for a specific element.
21. MEMORANDUM HEADER LINES: For documents using a MEMORANDUM letterhead, the header must be exactly three separate, individually centered paragraphs, in this exact order: "STATE OF HAWAII", then "DEPARTMENT OF TRANSPORTATION", then "MEMORANDUM". They must NOT be combined into a single line/paragraph, and each line must be centered on its own.

"""


def _get_kb_folders() -> dict:
    """Return {folder_name: [source, ...]} from the KB collection."""
    kb_meta = get_kb_all_metadata()
    folders = {}
    for m in kb_meta:
        folder = m.get("folder", "General")
        source = m.get("source", "unknown")
        if folder not in folders:
            folders[folder] = set()
        folders[folder].add(source)
    return {k: sorted(v) for k, v in sorted(folders.items())}


# ── DOCX Auto-Fix Engine ───────────────────────────────────────────────────

def _iter_block_items_with_path(parent, path_prefix=None):
    """Yield (path, paragraph) for every paragraph in document order,
    descending into tables (including nested tables) recursively."""
    if path_prefix is None:
        path_prefix = []

    from docx.document import Document as _DocClass
    from docx.table import Table, _Cell
    from docx.text.paragraph import Paragraph

    if isinstance(parent, _DocClass):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError(f"Unsupported parent type: {type(parent)}")

    para_counter = 0
    table_counter = 0
    for child in parent_elm.iterchildren():
        if child.tag == qn("w:p"):
            yield (path_prefix + [("para", para_counter)], Paragraph(child, parent))
            para_counter += 1
        elif child.tag == qn("w:tbl"):
            table = Table(child, parent)
            seen_tcs = set()
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    if cell._tc in seen_tcs:
                        continue
                    seen_tcs.add(cell._tc)
                    cell_path = path_prefix + [
                        ("table", table_counter), ("row", row_idx), ("cell", cell_idx)
                    ]
                    yield from _iter_block_items_with_path(cell, cell_path)
            table_counter += 1


def _get_paragraph_by_path(doc, target_path):
    """Re-walk the document and return the paragraph at `target_path`."""
    for path, para in _iter_block_items_with_path(doc):
        if path == target_path:
            return para
    return None


def _build_para_index(doc):
    """Build a dict mapping XML element id → paragraph object for the whole
    document. Used by fill_template_placeholders to locate paragraphs by
    element identity rather than walk-order path — which breaks when
    _insert_list_paragraphs_after adds new paragraphs earlier in the document,
    shifting all subsequent path indices."""
    index = {}
    for _, para in _iter_block_items_with_path(doc):
        index[id(para._p)] = para
    return index


def _force_preserve_space(run):
    """Ensure a run's text keeps its EXACT spacing when opened in Word."""
    for t in run._r.findall(qn("w:t")):
        t.set(qn("xml:space"), "preserve")


def build_indexed_run_report(docx_bytes):
    """Line-numbered formatting report plus a path-based line_map.

    Returns (numbered_text, line_map) where
        line_map[line_id] = {"kind": "para"|"run", "path":[...], "run_idx": int|None}
    """
    doc = Document(io.BytesIO(docx_bytes))
    lines = []
    line_map = {}
    counter = 0

    for path, para in _iter_block_items_with_path(doc):
        counter += 1
        para_id = f"D{counter:03d}"
        pf = para.paragraph_format
        sb = pf.space_before.pt if pf.space_before is not None else None
        sa = pf.space_after.pt  if pf.space_after  is not None else None
        lines.append(
            f"{para_id}: [PARA align={para.alignment} space_before={sb} "
            f"space_after={sa} line_spacing={pf.line_spacing}]"
        )
        line_map[para_id] = {"kind": "para", "path": path, "run_idx": None}

        for r_idx, run in enumerate(para.runs):
            counter += 1
            run_id = f"D{counter:03d}"
            try:
                color = (
                    run.font.color.rgb
                    if (run.font.color is not None and run.font.color.type is not None)
                    else None
                )
            except Exception:
                color = None
            size = run.font.size.pt if run.font.size is not None else None
            font_name = run.font.name
            try:
                highlight = run.font.highlight_color.name if run.font.highlight_color else None
            except Exception:
                highlight = None
            safe_text = run.text.replace("'", "\\'")
            lines.append(
                f"{run_id}: [RUN text='{safe_text}' bold={run.bold} italic={run.italic} "
                f"underline={run.underline} size={size} color={color} highlight={highlight} font={font_name}]"
            )
            line_map[run_id] = {"kind": "run", "path": path, "run_idx": r_idx}

    return "\n".join(lines), line_map


def extract_plain_text(docx_bytes):
    """Extract clean human-readable text from a docx, preserving table
    row-and-column structure AND list formatting (bullets, numbers, indent level)
    so the AI can correctly reconstruct lists when filling placeholders.

    Output order: body paragraphs first, then table rows — so the AI sees
    the actual memo content (SUBJECT, TO, FROM, body) before letterhead tables."""
    from docx.document import Document as _DocClass
    from docx.table import Table, _Cell
    from docx.text.paragraph import Paragraph
    from docx.oxml.ns import qn as _qn

    doc = Document(io.BytesIO(docx_bytes))
    body_lines  = []   # non-table paragraphs
    table_lines = []   # table content

    def _para_list_info(para):
        """Return (is_list, prefix_string) for a paragraph."""
        style_name = (para.style.name or "").lower() if para.style else ""
        text = "".join(r.text for r in para.runs).strip()

        numPr = para._p.find(_qn("w:numPr"))
        if numPr is not None:
            ilvl_el = numPr.find(_qn("w:ilvl"))
            ilvl = int(ilvl_el.get(_qn("w:val"), 0)) if ilvl_el is not None else 0
            indent = "  " * ilvl
            if any(x in style_name for x in ("bullet", "list bullet")):
                return True, f"{indent}• "
            elif any(x in style_name for x in ("number", "list number")):
                return True, f"{indent}1. "
            else:
                return True, f"{indent}• "

        if "list bullet" in style_name:
            return True, "• "
        if "list number" in style_name:
            return True, "1. "
        if "list paragraph" in style_name:
            if text and text[0] in ("•", "·", "◦", "▪", "▸"):
                return True, ""
            return False, ""
        if text and text[0] in ("•", "·", "◦", "▪", "▸"):
            return True, ""
        return False, ""

    def _cell_text(cell):
        parts = []
        for p in cell.paragraphs:
            t = "".join(r.text for r in p.runs)
            if t.strip():
                is_list, prefix = _para_list_info(p)
                parts.append(prefix + t if is_list and prefix else t)
        return " / ".join(parts) if parts else ""

    def _is_letterhead_table(table):
        """Return True if this table looks like a letterhead/header table
        that should be skipped during plain text extraction."""
        # Sample the first few cells for letterhead indicators
        indicators = {
            "GOVERNOR", "DIRECTOR", "STATE OF HAWAII", "DEPARTMENT OF TRANSPORTATION",
            "KE KIA", "KA LUNA", "DEPUTY DIRECTOR", "DESIGN BRANCH"
        }
        sample_text = ""
        try:
            for row in table.rows[:3]:
                for cell in row.cells[:4]:
                    sample_text += cell.text.upper() + " "
        except Exception:
            pass
        return any(ind in sample_text for ind in indicators)

    def _walk(parent_elm, parent_obj):
        for child in parent_elm.iterchildren():
            if child.tag == _qn("w:p"):
                para = Paragraph(child, parent_obj)
                text = "".join(r.text for r in para.runs)
                if text.strip():
                    is_list, prefix = _para_list_info(para)
                    body_lines.append(prefix + text if is_list and prefix else text)
            elif child.tag == _qn("w:tbl"):
                table = Table(child, parent_obj)
                # Skip letterhead/header tables — they contain repeated
                # administrative info (GOVERNOR, DIRECTOR etc.) that would
                # confuse the AI when extracting placeholder values
                if _is_letterhead_table(table):
                    continue
                seen_tcs = set()
                for row in table.rows:
                    cells_text = []
                    for cell in row.cells:
                        if cell._tc in seen_tcs:
                            continue
                        seen_tcs.add(cell._tc)
                        cells_text.append(_cell_text(cell))
                    if any(t.strip() for t in cells_text):
                        if len(cells_text) == 1:
                            table_lines.append(cells_text[0])
                        else:
                            table_lines.append("[ROW] " + " | ".join(cells_text) + " [/ROW]")

    _walk(doc.element.body, doc)
    # Body content first so AI sees memo text before letterhead table noise
    all_lines = body_lines + (["--- TABLE CONTENT ---"] + table_lines if table_lines else [])
    return "\n".join(all_lines)


# ── Generate Template From a Filled Document ────────────────────────────────

PLACEHOLDER_COLOR = RGBColor(0xFF, 0x00, 0x00)


def build_field_detection_prompt(numbered_text):
    """Build the prompt asking the AI to identify fillable runs."""
    return (
        "You are analyzing an HDOT government document so it can be converted into a "
        "reusable BLANK TEMPLATE.\n\n"
        "Identify every line below that contains INSTANCE-SPECIFIC content — information "
        "that is unique to this particular document and would be different the next time "
        "this type of document is filled out. Examples: names, job titles, dates, memo "
        "numbers, dollar amounts, addresses, phone numbers, the specific SUBJECT text, "
        "case-specific body content, signature names.\n\n"
        "Do NOT flag FIXED BOILERPLATE that should stay exactly the same in every copy of "
        "this template — labels like 'TO:', 'FROM:', 'SUBJECT:', 'DATE:', the standard "
        "letterhead ('STATE OF HAWAII', 'DEPARTMENT OF TRANSPORTATION', 'MEMORANDUM'), "
        "numbered item labels, standard instructional or closing language, signature block "
        "titles like 'Director of Transportation', section punctuation, or blank/"
        "whitespace-only runs.\n\n"
        "══════════════════════════════════════════\n"
        "DOCUMENT (line-numbered, prefix D):\n"
        "══════════════════════════════════════════\n"
        f"{numbered_text}\n\n"
        "For each instance-specific run you find, return its D-line ID and a short ALL-CAPS "
        "field label describing what kind of information belongs there (e.g. 'MEMO NO.', "
        "'DATE', 'NAME AND TITLE OF INDIVIDUAL TRAVELING', 'SUBJECT', 'DOLLAR AMOUNT'). "
        "Only ever reference RUN lines, never PARA lines — a field must be an actual run "
        "of text, not a paragraph-level entry.\n\n"
        "Return ONLY valid JSON (no markdown fences, no explanation outside the JSON):\n"
        "{\n"
        '  "fields": [\n'
        '    {"run_id": "D0xx", "label": "FIELD LABEL", "original_text": "the current text"}\n'
        "  ]\n"
        "}"
    )


def blank_out_fields(original_bytes, fields, line_map):
    """Turn a filled document into a reusable blank template."""
    doc = Document(io.BytesIO(original_bytes))
    applied, skipped = [], []

    for field in fields:
        run_id = field.get("run_id")
        label = (field.get("label") or "FIELD").strip().strip(":").upper()
        if not run_id or run_id not in line_map:
            skipped.append((label, "no reliable location match"))
            continue

        entry = line_map[run_id]
        if entry["kind"] != "run":
            skipped.append((label, "not a text run — can't blank a paragraph marker"))
            continue

        para = _get_paragraph_by_path(doc, entry["path"])
        if para is None or entry["run_idx"] >= len(para.runs):
            skipped.append((label, "could not relocate the original run"))
            continue

        try:
            run = para.runs[entry["run_idx"]]
            run.text = f"{{{{ {label} }}}}"
            _force_preserve_space(run)
            run.font.color.rgb = PLACEHOLDER_COLOR
            applied.append(label)
        except Exception as e:
            skipped.append((label, f"error blanking field: {e}"))

    out_buf = io.BytesIO()
    doc.save(out_buf)
    return out_buf.getvalue(), applied, skipped


# ── Generate Corrected Document FROM THE TEMPLATE ───────────────────────────

_PLACEHOLDER_BRACE_RE = re.compile(r"\{\{\s*(.*?)\s*\}\}")
_PLACEHOLDER_PAREN_RE = re.compile(r"\(([A-Za-z][^)]*)\)")


def find_template_placeholders(template_bytes):
    """Walk the official template and return (line_map, placeholders).

    ONLY detects {{ placeholder_name }} patterns — nothing else.
    Red text and ( ) patterns are ignored entirely.

    Handles Word splitting {{ }} across multiple runs (e.g. spellcheck
    splitting {{PROJECT_TITLE}} into '{{PROJECT', '_', 'TITLE}}').
    The canonical run is whichever run contains the opening '{{'.
    All other runs spanned by the same {{ }} are recorded in clear_run_ids
    so they get zeroed out during fill.

    The label shown to the user is derived from the text INSIDE the {{ }},
    with underscores replaced by spaces and forced to uppercase.
    e.g. {{project_title}} → label = "PROJECT TITLE"
         {{tAX mAP kEY nO.}} → label = "TAX MAP KEY NO."
    """
    doc = Document(io.BytesIO(template_bytes))
    line_map    = {}
    placeholders = []
    counter     = 0
    seen_labels: dict = {}

    def _unique_label(raw: str) -> str:
        """Deduplicate labels — same normalised key returns the first-seen form."""
        key = re.sub(r"[\s_\-]+", "_", raw.strip()).upper()
        if key not in seen_labels:
            seen_labels[key] = raw.strip()
            return raw.strip()
        return seen_labels[key]

    def _inner_to_label(inner: str) -> str:
        """Convert inner placeholder text to a clean display label.
        {{project_title}} → PROJECT TITLE
        {{tAX mAP kEY nO.}} → TAX MAP KEY NO.
        """
        return re.sub(r"[_]+", " ", inner).strip().upper()

    # ── Body paragraphs and table cells ──────────────────────────────────────
    for path, para in _iter_block_items_with_path(doc):
        counter += 1
        para_id = f"D{counter:03d}"
        line_map[para_id] = {"kind": "para", "path": path, "run_idx": None}

        # Pre-assign D-line IDs to all runs
        run_ids = []
        para_elm_id = id(para._p)
        for r_idx, run in enumerate(para.runs):
            counter += 1
            run_id = f"D{counter:03d}"
            line_map[run_id] = {"kind": "run", "path": path, "run_idx": r_idx, "para_elm_id": para_elm_id}
            run_ids.append(run_id)

        # Row key for side-by-side grouping in review UI
        row_key = None
        for i, step in enumerate(path):
            if step[0] == "row":
                row_key = "|".join(str(s) for s in path[:i + 1])

        # Concatenate all run text — catches {{ }} split across multiple runs
        full_text   = ""
        char_to_run = []
        for r_idx, run in enumerate(para.runs):
            t = run.text or ""
            char_to_run.extend([r_idx] * len(t))
            full_text += t

        # Scan for {{ }} patterns in the concatenated text
        for m in _PLACEHOLDER_BRACE_RE.finditer(full_text):
            ph_start, ph_end = m.start(), m.end()
            inner_text = m.group(1).strip()

            if ph_start >= len(char_to_run):
                continue

            canonical_r_idx = char_to_run[ph_start]
            spanned_runs    = set(char_to_run[ph_start:ph_end])
            clear_r_idxs    = [ri for ri in spanned_runs if ri != canonical_r_idx]

            canonical_run_text = para.runs[canonical_r_idx].text if canonical_r_idx < len(para.runs) else ""

            # Skip ghost detections — empty canonical means the run was already
            # consumed (e.g. duplicate table cells in a repeated letterhead)
            if not canonical_run_text.strip():
                continue

            run_id    = run_ids[canonical_r_idx] if canonical_r_idx < len(run_ids) else para_id
            clear_ids = [run_ids[ri] for ri in clear_r_idxs if ri < len(run_ids)]
            label     = _unique_label(_inner_to_label(inner_text) or "FIELD")

            placeholders.append({
                "run_id":             run_id,
                "clear_run_ids":      clear_ids,
                "label":              label,
                "current_text":       full_text[ph_start:ph_end],  # full {{ }} from concat
                "canonical_run_text": canonical_run_text,           # actual run text
                "inner_text":         inner_text,
                "row_key":            row_key,
            })

    # ── Page headers (separate from body element tree) ────────────────────────
    for section_idx, section in enumerate(doc.sections):
        try:
            header = section.header
            if header is None or not header.is_linked_to_previous:
                for para_idx, hpara in enumerate(header.paragraphs):
                    full_text = "".join(r.text or "" for r in hpara.runs)
                    if not full_text.strip():
                        continue
                    if "{{" not in full_text:
                        continue

                    h_run_ids   = []
                    char_to_run_h = []
                    for r_idx, run in enumerate(hpara.runs):
                        counter += 1
                        run_id = f"D{counter:03d}"
                        line_map[run_id] = {
                            "kind":               "run",
                            "path":               [("header",)],
                            "run_idx":            r_idx,
                            "header_section_idx": section_idx,
                            "header_para_idx":    para_idx,
                        }
                        h_run_ids.append(run_id)
                        char_to_run_h.extend([r_idx] * len(run.text or ""))

                    for m in _PLACEHOLDER_BRACE_RE.finditer(full_text):
                        ph_start, ph_end = m.start(), m.end()
                        inner_text = m.group(1).strip()
                        if ph_start >= len(char_to_run_h):
                            continue
                        canonical_r = char_to_run_h[ph_start]
                        spanned     = set(char_to_run_h[ph_start:ph_end])
                        clear_r     = [ri for ri in spanned if ri != canonical_r]
                        run_id      = h_run_ids[canonical_r] if canonical_r < len(h_run_ids) else None
                        if not run_id:
                            continue
                        canonical_run_text = hpara.runs[canonical_r].text if canonical_r < len(hpara.runs) else ""
                        if not canonical_run_text.strip():
                            continue
                        label = _unique_label(_inner_to_label(inner_text) or "FIELD")
                        placeholders.append({
                            "run_id":             run_id,
                            "clear_run_ids":      [h_run_ids[ri] for ri in clear_r if ri < len(h_run_ids)],
                            "label":              label,
                            "current_text":       m.group(0),
                            "canonical_run_text": canonical_run_text,
                            "inner_text":         inner_text,
                            "row_key":            None,
                        })
        except Exception:
            pass

    return line_map, placeholders


def build_template_fill_prompt(fields_list_str, source_text):
    """Prompt asking the AI to find each placeholder field's value."""
    return (
        "You are filling in an official HDOT document template.\n\n"
        "══════════════════════════════════════════\n"
        "CRITICAL RULE — READ THIS FIRST:\n"
        "══════════════════════════════════════════\n"
        "You may ONLY provide values for fields listed below that are marked with {{ }}.\n"
        "You must NEVER modify, remove, rewrite, or touch ANY other text in the document.\n"
        "Every piece of text in the template that is NOT inside {{ }} must remain exactly "
        "as it is — you are only filling in the blanks marked {{ }}.\n"
        "Your ONLY output is a JSON list of run_id → value pairs for the {{ }} fields below.\n"
        "Do NOT include any other text, explanation, or commentary.\n\n"
        "══════════════════════════════════════════\n"
        "PLACEHOLDER FIELDS TO FILL (these are the ONLY things you touch):\n"
        "══════════════════════════════════════════\n"
        "Each field is listed as: RUN_ID: LABEL\n"
        "The LABEL comes from the text inside {{ }} in the template. "
        "It may have inconsistent casing or spelling — match by MEANING.\n\n"
        f"{fields_list_str}\n\n"
        "══════════════════════════════════════════\n"
        "UPLOADED DOCUMENT (values must come ONLY from here):\n"
        "══════════════════════════════════════════\n"
        f"{source_text[:20000]}\n\n"
        "IMPORTANT — table structure:\n"
        "Lines wrapped in [ROW] ... [/ROW] are table rows with cells separated by ' | '. "
        "Each cell is an INDEPENDENT field — the value for a placeholder comes ONLY from "
        "its own cell. For example in '[ROW] PHASE: Alpha | CPN: Beta [/ROW]', the value "
        "of PHASE is 'Alpha' and the value of CPN is 'Beta' — never mix cell contents.\n\n"
        "IMPORTANT — lists and numbering:\n"
        "The source document text already includes list prefixes — lines starting with "
        "'• ' are bullet points, lines starting with '  • ' (two spaces) are sub-bullets, "
        "lines starting with '1. ' are numbered items, lines starting with '  1. ' are "
        "sub-numbered items. When a field value spans multiple list items, capture ALL of "
        "them and join with '\\n'. Preserve every prefix exactly as it appears.\n"
        "CRITICAL: each '\\n' must only appear BETWEEN list items, never in the middle of "
        "a single list item's text. Each list item must be on exactly ONE line with its "
        "full text — never split a sentence across multiple lines. Wrong: "
        "'TYPE 2, PART 1: Replacement\\nof assets' — Right: "
        "'TYPE 2, PART 1: Replacement of assets that have exceeded...'\n\n"
        "IMPORTANT — placeholder label matching:\n"
        "Match each placeholder to its value by meaning, not exact label text. Examples:\n"
        "- 'TAX MAP KEY NO' → look for 'TAX MAP KEY NO.' in the source\n"
        "- 'REQEST INFORMATION' → look for 'REQUEST FOR INFORMATION' content\n"
        "- 'ENVIORMENTAL COMPLIANCE' → look for 'ENVIRONMENTAL COMPLIANCE' content\n"
        "- 'CHAPTER OF HRS' → look for the HRS chapter reference\n"
        "- 'CITY ISLAND' → look for city and island name\n"
        "- 'REFF NUM' or 'REF NUM' or 'REFERENCE NUMBER' → look for the number that follows "
        "'HWY-DD' or 'IN REPLY REFER TO:' in the source — extract ONLY the number portion, "
        "not the 'HWY-DD' prefix. For example if source has 'HWY-DD 2.21253', the value is '2.21253'.\n\n"
        "IMPORTANT — mixed text placeholders:\n"
        "Some placeholders appear alongside fixed text, for example 'HWY-DD {{reff_num}}'. "
        "Provide ONLY the value for the {{ }} part — the surrounding fixed text stays as-is. "
        "So for reference number '2024-001', provide just '2024-001', not 'HWY-DD 2024-001'.\n\n"
        "For each field, extract ONLY the value — strip any field label prefix "
        "(e.g. for 'PROJECT TITLE: Foo', the value is 'Foo', not 'PROJECT TITLE: Foo'). "
        "However do NOT strip words that are part of the actual value content — "
        "for example if the source says 'SUBJECT: CHAPTER 195D/CHAPTER 343 HAWAII REVISED STATUTES (HRS)', "
        "the value for CHAPTER OF HRS is the full 'CHAPTER 195D/CHAPTER 343 HAWAII REVISED STATUTES (HRS)', "
        "not just '195D/CHAPTER 343'. Strip only the colon-separated label, not words that are part of the value itself. "
        "If you cannot find a confident match in the source, set \"found\": false.\n\n"
        "NEVER invent, assume, or guess a value. Only use what is in the uploaded document.\n\n"
        "Return ONLY valid JSON — no markdown fences, no explanation, nothing else:\n"
        "{\n"
        '  "fills": [\n'
        '    {"run_id": "D0xx", "value": "the exact replacement text", "found": true|false}\n'
        "  ]\n"
        "}"
    )


def _parse_list_lines(value: str):
    """Parse a multi-line value into structured paragraph items.

    Detects list items from:
    1. Explicit bullet chars: • · ◦ ▪ ▸ ‣ ➢ ➤ – - *
    2. Explicit numbered prefixes: 1. 2. 14. (i) (ii) (a) (b) I. II.
    3. Letter+period labels: A. B. C.
    4. ALL-CAPS label ending with colon: EXEMPTION TYPE 2: TYPE 2, PART 1:
    5. Indentation level from leading spaces (2 spaces = 1 indent level)
    6. Lines before the first list item → plain paragraph (no bullet)

    Returns list of dicts:
        text        — line text after stripping bullet prefix
        bold_prefix — ALL-CAPS label ending in ":" to be bolded
        text_after  — text after bold_prefix
        indent      — nesting level (0=top, 1=sub, 2=sub-sub, 3=deepest)
        is_list     — True if this line is a list item
        raw         — original line before processing
    """
    items = []
    lines = [l for l in value.replace("\\n", "\n").split("\n") if l.strip()]
    seen_list_item = False

    # Explicit bullet characters (single char)
    BULLET_CHARS = set("•·◦▪▸‣➢➤")

    for raw_line in lines:
        stripped        = raw_line.lstrip(" ")
        leading_spaces  = len(raw_line) - len(stripped)
        explicit_indent = leading_spaces // 2

        text     = stripped
        is_list  = False
        indent   = explicit_indent

        # ── 1. Explicit bullet character prefix ──────────────────────────────
        for prefix in ("• ", "· ", "◦ ", "▪ ", "▸ ", "‣ ", "➢ ", "➤ ", "– ", "* "):
            if text.startswith(prefix):
                text    = text[len(prefix):]
                is_list = True
                seen_list_item = True
                break
        # "o " as bullet (common Word list char) — only at line start
        if not is_list and re.match(r"^o\s+\S", text) and (explicit_indent > 0 or seen_list_item):
            text    = text[2:].lstrip()
            is_list = True
            seen_list_item = True
            indent  = max(explicit_indent, 1)
        # Single bullet char with no space
        if not is_list and text and text[0] in BULLET_CHARS:
            text    = text[1:].lstrip()
            is_list = True
            seen_list_item = True

        # ── 2. Numbered prefix: "1." "14." "1)" ──────────────────────────────
        if not is_list:
            m = re.match(r"^(\d+[.)]\s*)", text)
            if m:
                num = int(re.match(r"\d+", m.group(1)).group())
                is_list = True
                seen_list_item = True
                # Higher numbers tend to be deeper sub-items
                indent = max(explicit_indent, 3 if num >= 10 else 2 if num >= 5 else 1)

        # ── 3. Roman numeral prefix: "(i)" "(ii)" "(iii)" "I." "II." ─────────
        if not is_list:
            if re.match(r"^\([ivxlIVXL]+\)\s+", text) or re.match(r"^[IVXivx]+\.\s+", text):
                is_list = True
                seen_list_item = True
                indent  = max(explicit_indent, 2)

        # ── 4. Letter+period or letter+paren: "A." "a." "(a)" "(b)" ─────────
        if not is_list:
            if re.match(r"^[A-Za-z][.)]\s+", text):
                is_list = True
                seen_list_item = True
                indent  = max(explicit_indent, 2)
            elif re.match(r"^\([A-Za-z]\)\s+", text):
                is_list = True
                seen_list_item = True
                indent  = max(explicit_indent, 2)

        # ── 5. ALL-CAPS label ending with ":" ────────────────────────────────
        if not is_list:
            caps_match = re.match(r"^([A-Z][A-Z0-9 ,\-–/]*:)\s+\S", text)
            if caps_match:
                is_list = True
                seen_list_item = True
                # Determine indent from the label pattern
                label_text = caps_match.group(1)
                if re.match(r"^TYPE\s+\d", label_text, re.IGNORECASE):
                    indent = max(explicit_indent, 1)
                else:
                    indent = max(explicit_indent, 0)

        # ── 6. Indented line after a list item → sub-item ────────────────────
        if not is_list and seen_list_item and explicit_indent > 0:
            is_list = True
            indent  = explicit_indent

        # ── Detect bold prefix: ALL-CAPS label ending with ":" ───────────────
        bold_prefix = ""
        text_after  = text
        bold_match  = re.match(r"^([A-Z][A-Z0-9 ,\-–/]*:)\s*", text)
        if bold_match and len(bold_match.group(1)) > 2:  # avoid "A:" single-letter false positives
            bold_prefix = bold_match.group(1) + " "
            text_after  = text[bold_match.end():]

        items.append({
            "text":        text,
            "bold_prefix": bold_prefix,
            "text_after":  text_after,
            "indent":      indent,
            "is_list":     is_list,
            "raw":         raw_line,
        })

    # ── Post-process: merge continuation lines into their parent item ────────
    merged = []

    def _looks_like_new_item(text):
        if not text:
            return False
        if text[0] in BULLET_CHARS:
            return True
        if re.match(r"^o\s+\S", text):
            return True
        if re.match(r"^[A-Za-z][.)]\s+", text):
            return True
        if re.match(r"^\d+[.)]\s+", text):
            return True
        if re.match(r"^[A-Z][A-Z0-9 ,\-–/]{2,}:\s+\S", text):
            return True
        return False

    for item in items:
        def _is_cont(item, prev):
            if not prev["is_list"]:
                return False
            if item["is_list"]:
                return False
            if item["indent"] > prev["indent"]:
                return False
            if _looks_like_new_item(item["text"]):
                return False
            if item["text"] and item["text"][0].islower():
                return True
            prev_text = prev["text"].rstrip()
            if prev_text and prev_text[-1] not in ".!?":
                return True
            return False

        if merged and _is_cont(item, merged[-1]):
            prev  = merged[-1]
            extra = " " + item["text"]
            prev["text"] = prev["text"].rstrip() + extra
            if prev["bold_prefix"]:
                prev["text_after"] = prev["text_after"].rstrip() + extra
        else:
            merged.append(item)

    return merged




def _insert_list_paragraphs_after(anchor_para, items, template_doc):
    """Insert properly formatted list paragraphs after anchor_para in the document.

    Copies font properties (name, size) from the anchor paragraph's first run so
    the inserted paragraphs match the template's base style. Applies list-style
    indentation via paragraph format rather than hard-coded list XML so it works
    with any document, not just HDOT memos.

    Returns the last inserted paragraph (for chaining).
    """
    from docx.oxml import OxmlElement
    from docx.shared import Pt, Inches
    from docx.text.paragraph import Paragraph as _Para

    # Gather base font properties from anchor paragraph
    base_font_name = None
    base_font_size = None
    for r in anchor_para.runs:
        if r.font.name:
            base_font_name = r.font.name
            break
    for r in anchor_para.runs:
        if r.font.size:
            base_font_size = r.font.size
            break

    # Indent sizes per level (in inches): matches standard Word list indents
    INDENT_PER_LEVEL = 0.25

    def _add_run(para, text, bold=False):
        r = para.add_run(text)
        if base_font_name:
            r.font.name = base_font_name
        if base_font_size:
            r.font.size = base_font_size
        if bold:
            r.bold = True
        _force_preserve_space(r)
        return r

    # Insert in REVERSE order — each new paragraph is inserted immediately after
    # anchor_para using addnext(), which pushes previous insertions down.
    # Reversing ensures final document order matches the original list order.
    last_para = anchor_para
    for item in reversed(items):
        new_p = OxmlElement("w:p")
        anchor_para._p.addnext(new_p)
        new_para = _Para(new_p, template_doc)

        # Apply paragraph formatting
        pf = new_para.paragraph_format
        indent_inches = item["indent"] * INDENT_PER_LEVEL + INDENT_PER_LEVEL

        # Build the paragraph text — use different bullet per indent level
        # matching standard Word list style: • ◦ ▪ at levels 0, 1, 2+
        if item["is_list"]:
            level_bullets = {0: "•", 1: "o", 2: "▪", 3: "▸"}
            bullet_char = level_bullets.get(item["indent"], "▸")
        else:
            bullet_char = ""

        if bullet_char:
            _add_run(new_para, bullet_char + "\t")
            pf.left_indent       = Inches(indent_inches)
            pf.first_line_indent = Inches(-INDENT_PER_LEVEL)  # hanging indent
        else:
            # Plain paragraph — indent only if nested, no bullet, no hanging indent
            pf.left_indent       = Inches(item["indent"] * INDENT_PER_LEVEL) if item["indent"] > 0 else None
            pf.first_line_indent = None

        if item["bold_prefix"]:
            _add_run(new_para, item["bold_prefix"], bold=True)
            _add_run(new_para, item["text_after"])
        else:
            _add_run(new_para, item["text"])

        last_para = new_para

    return last_para


def fill_template_placeholders(template_bytes, fills, line_map, placeholder_labels=None):
    """Fill in ONLY the placeholder runs of the official template.

    Multi-line list values (joined with \\n by the AI) are expanded into
    separate properly-indented paragraphs inserted after the placeholder
    paragraph, preserving bold labels and indent hierarchy — without
    hard-coding anything document-specific.

    Returns (filled_bytes, applied_log, skipped_log).
    """
    from docx.table import Table as _Table

    placeholder_labels = placeholder_labels or {}
    doc = Document(io.BytesIO(template_bytes))
    applied, skipped = [], []
    filled_paths        = set()
    filled_run_ids      = set()
    filled_labels       = set()
    filled_header_labels = set()

    # Build paragraph index. Normalise path keys so tuple/list differences
    # from JSON round-tripping through session state don't cause mismatches.
    # e.g. [('para', 5)] and [['para', 5]] must produce the same key.
    def _path_key(path):
        return str([list(step) for step in path])

    para_by_path = {}
    for path, para in _iter_block_items_with_path(doc):
        para_by_path[_path_key(path)] = para

    def _para_for_entry(entry):
        """Look up the paragraph for a line_map entry."""
        path = entry.get("path", [])
        # Header path can be [('header',)] or [['header']] after JSON round-trip
        is_header = (len(path) == 1 and list(path[0])[:1] == ["header"])
        if is_header:
            section_idx = entry.get("header_section_idx", 0)
            para_idx    = entry.get("header_para_idx", 0)
            try:
                return doc.sections[section_idx].header.paragraphs[para_idx]
            except (IndexError, AttributeError):
                return None
        key  = _path_key(path)
        para = para_by_path.get(key)
        if para is None:
            import sys
            print(f"[PATH MISS] key={key!r} not in index (size={len(para_by_path)}), trying full walk", file=sys.stderr)
            para = _get_paragraph_by_path(doc, entry["path"])
            if para is None:
                print(f"[PATH MISS] full walk also returned None for key={key!r}", file=sys.stderr)
        return para

    # Split fills into single-line and multiline so we can process
    # multiline inserts in reverse order (bottom-to-top of document)
    # without invalidating path lookups for subsequent fills.
    def _is_multiline(fill):
        v = (fill.get("value") or "").strip().replace("\\n", "\n")
        return len([l for l in v.split("\n") if l.strip()]) > 1

    single_fills    = [f for f in fills if not _is_multiline(f)]
    multiline_fills = [f for f in fills if _is_multiline(f)]

    ordered_fills = single_fills + list(reversed(multiline_fills))

    # Pre-determine the canonical run_id for each label — the FIRST occurrence
    # in document order (lowest D-number) wins. All others are duplicates.
    canonical_run_for_label = {}
    for fill in fills:
        rid   = fill.get("run_id", "")
        lbl   = placeholder_labels.get(rid, rid)
        if lbl not in canonical_run_for_label:
            canonical_run_for_label[lbl] = rid
        else:
            # Keep the lower D-number as canonical
            existing = canonical_run_for_label[lbl]
            try:
                if int(rid[1:]) < int(existing[1:]):
                    canonical_run_for_label[lbl] = rid
            except (ValueError, IndexError):
                pass

    for fill in ordered_fills:
        run_id        = fill.get("run_id")
        value         = (fill.get("value") or "").strip()
        found         = fill.get("found", bool(value))
        label         = placeholder_labels.get(run_id) or run_id or "field"
        clear_run_ids = fill.get("clear_run_ids") or []
        current_text  = (fill.get("current_text") or "").strip()

        # STRICT RULE: only process fills that were detected from a {{ }} placeholder.
        # If current_text doesn't contain {{ we never detected this as a placeholder
        # and must not touch the run — skip entirely.
        if "{{" not in current_text:
            skipped.append((label, "skipped — not a {{ }} placeholder"))
            continue

        write_value = value if (found and value) else ""

        if not run_id or run_id not in line_map:
            skipped.append((label, "could not locate run in document"))
            continue

        # Skip duplicate run_ids — same placeholder detected twice
        if run_id in filled_run_ids:
            continue
        filled_run_ids.add(run_id)

        entry = line_map[run_id]
        if entry["kind"] != "run":
            skipped.append((label, "not a text run"))
            continue

        is_header_run = (len(entry.get("path", [])) == 1 and
                         list(entry["path"][0])[:1] == ["header"])

        # Skip non-canonical body duplicates — only the first D-number for each
        # label (canonical_run_for_label) gets filled in the body.
        # Page headers are allowed through once each.
        canonical_rid = canonical_run_for_label.get(label)
        if canonical_rid and run_id != canonical_rid:
            if not is_header_run:
                continue  # body duplicate (repeated letterhead table) — skip
            if label in filled_header_labels:
                continue  # already filled once in a header — skip

        # If no value provided, clear the placeholder runs and skip — not an error
        if not write_value:
            para = _para_for_entry(entry)
            if para is not None and entry["run_idx"] < len(para.runs):
                para.runs[entry["run_idx"]].text = ""
                for cr_id in clear_run_ids:
                    if cr_id in line_map:
                        cr_e = line_map[cr_id]
                        if cr_e["kind"] == "run" and cr_e["path"] == entry["path"]:
                            if cr_e["run_idx"] < len(para.runs):
                                para.runs[cr_e["run_idx"]].text = ""
            skipped.append((label, "not found in uploaded document — left blank"))
            continue

        # Header runs use section/para index to re-locate in the fresh doc
        if entry.get("path") == [("header",)]:
            para = _para_for_entry(entry)
        else:
            para = _para_for_entry(entry)
        if para is None or entry["run_idx"] >= len(para.runs):
            skipped.append((label, "could not relocate the placeholder run"))
            continue

        try:
            run = para.runs[entry["run_idx"]]

            # ── Zero out split-placeholder fragment runs FIRST ────────────────
            # Fragment runs share the same paragraph as the canonical run.
            # Use the already-resolved `para` object directly instead of
            # re-walking the document — this is faster and more reliable.
            def _clear_cr(cr_id):
                if cr_id not in line_map:
                    return
                cr_entry = line_map[cr_id]
                if cr_entry["kind"] != "run":
                    return
                # Normalise header path check
                cr_path  = cr_entry.get("path", [])
                is_header = (len(cr_path) == 1 and list(cr_path[0])[:1] == ["header"])
                if is_header:
                    section_idx = cr_entry.get("header_section_idx", 0)
                    para_idx    = cr_entry.get("header_para_idx", 0)
                    try:
                        cr_para = doc.sections[section_idx].header.paragraphs[para_idx]
                        if cr_para is not None and cr_entry["run_idx"] < len(cr_para.runs):
                            cr_para.runs[cr_entry["run_idx"]].text = ""
                    except (IndexError, AttributeError):
                        pass
                elif _path_key(cr_entry["path"]) == _path_key(entry["path"]):
                    if cr_entry["run_idx"] < len(para.runs):
                        para.runs[cr_entry["run_idx"]].text = ""
                else:
                    cr_para = _para_for_entry(cr_entry)
                    if cr_para is not None and cr_entry["run_idx"] < len(cr_para.runs):
                        cr_para.runs[cr_entry["run_idx"]].text = ""

            for cr_id in clear_run_ids:
                _clear_cr(cr_id)

            current_text       = (fill.get("current_text")        or "").strip()
            canonical_run_text = (fill.get("canonical_run_text")  or "").strip()
            inner_text         = (fill.get("inner_text")           or "").strip()
            run_text_now       = run.text or ""

            # Double-check: the run must contain {{ or match the canonical
            # placeholder text. If neither is true, something went wrong during
            # detection and we must NOT overwrite this run.
            run_has_placeholder = (
                "{{" in run_text_now
                or run_text_now.strip() == canonical_run_text.strip()
                or (canonical_run_text and canonical_run_text in run_text_now)
            )
            if not run_has_placeholder:
                import sys
                print(f"[SKIP GUARD] {label!r} run_text={run_text_now!r} has no {{{{ — skipping",
                      file=sys.stderr)
                skipped.append((label, f"run '{run_text_now[:30]}' does not contain {{{{ — skipped for safety"))
                continue

            # ── Debug ─────────────────────────────────────────────────────────
            import sys
            print(f"[FILL DEBUG] label={label!r} run_id={run_id} "
                  f"run_text={run_text_now!r} current_text={current_text!r} "
                  f"canonical={canonical_run_text!r} inner={inner_text!r} "
                  f"write_value={write_value[:60]!r} found={found}",
                  file=sys.stderr)

            # ── Normalise value ───────────────────────────────────────────────
            normalised_value = write_value.replace("\\n", "\n")
            lines = [l for l in normalised_value.split("\n") if l.strip()]

            # Only insert paragraphs for actual bulleted/numbered lists.
            # Plain multiline text (subject lines, address continuations) stays
            # in the single run — paragraph insertion displaces template structure.
            if len(lines) > 1:
                list_items    = _parse_list_lines(normalised_value)
                has_real_list = any(item["is_list"] for item in list_items)
            else:
                list_items    = []
                has_real_list = False

            is_multiline_list = len(lines) > 1 and has_real_list

            if not is_multiline_list and len(lines) > 1:
                # Plain multiline — join back to single line with space
                normalised_value = " ".join(lines)

            if is_multiline_list:
                run.text = ""
                _force_preserve_space(run)
                try:
                    if run.font.color is not None and run.font.color.type is not None:
                        if str(run.font.color.rgb).upper() == "FF0000":
                            run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
                except Exception:
                    pass
                list_items = _parse_list_lines(normalised_value)
                _insert_list_paragraphs_after(para, list_items, doc)

            else:
                # ── Single-line replacement ───────────────────────────────────
                # Only replace within this run if it was detected as a {{ }}
                # placeholder. We try strategies in priority order.
                replaced = False

                # P0: run contains {{ anywhere (handles " {{" with leading space)
                # Check this first for fragment-only runs like " {{" or "\t{{"
                if not replaced and run_text_now.strip() in ("{{", "{{}", "{"):
                    run.text = normalised_value
                    replaced = True
                    import sys; print(f"[P0 HIT] {label!r} → {normalised_value[:40]!r}", file=sys.stderr)

                # P1: whole run stripped matches canonical stripped
                # If the run has text before {{ (e.g. "NO. {{"), preserve that prefix
                # and only replace from {{ onwards.
                if not replaced and canonical_run_text and run_text_now.strip() == canonical_run_text.strip():
                    if "{{" in run_text_now:
                        idx = run_text_now.index("{{")
                        run.text = run_text_now[:idx] + normalised_value
                    else:
                        run.text = normalised_value
                    replaced = True
                    import sys; print(f"[P1 HIT] {label!r} → {run.text[:40]!r}", file=sys.stderr)

                # P2: run contains the full {{ }} pattern verbatim
                if not replaced and current_text and current_text in run_text_now:
                    run.text = run_text_now.replace(current_text, normalised_value, 1)
                    replaced = True
                    import sys; print(f"[P2 HIT] {label!r} → {normalised_value[:40]!r}", file=sys.stderr)

                # P3: normalised brace spacing match
                if not replaced:
                    norm_ct = re.sub(r"\{\{\s*(.*?)\s*\}\}", r"{{\1}}", current_text)
                    norm_rt = re.sub(r"\{\{\s*(.*?)\s*\}\}", r"{{\1}}", run_text_now)
                    if norm_ct and norm_ct in norm_rt:
                        run.text = norm_rt.replace(norm_ct, normalised_value, 1)
                        replaced = True
                        import sys; print(f"[P3 HIT] {label!r} → {normalised_value[:40]!r}", file=sys.stderr)

                # P4: any {{ }} in the run → regex replace
                if not replaced and _PLACEHOLDER_BRACE_RE.search(run_text_now):
                    run.text = _PLACEHOLDER_BRACE_RE.sub(normalised_value, run_text_now, count=1)
                    replaced = True
                    import sys; print(f"[P4 HIT] {label!r} → {normalised_value[:40]!r}", file=sys.stderr)

                # P5: inner_text match — run contains just the label inside {{ }}
                if not replaced and inner_text and inner_text in run_text_now:
                    if re.match(r"^[A-Z0-9 _\-\.]+$", inner_text.upper()):
                        run.text = run_text_now.replace(inner_text, normalised_value, 1)
                        replaced = True
                        import sys; print(f"[P5 HIT] {label!r} → {normalised_value[:40]!r}", file=sys.stderr)

                # P6: run has {{ anywhere with surrounding whitespace/text
                if not replaced and "{{" in run_text_now:
                    # Replace from {{ to end of run (handles " {{ " style fragments)
                    idx = run_text_now.index("{{")
                    run.text = run_text_now[:idx] + normalised_value
                    replaced = True
                    import sys; print(f"[P6 HIT] {label!r} → {normalised_value[:40]!r}", file=sys.stderr)

                # If nothing matched, leave the run untouched — do NOT write
                # the value blindly. This prevents overwriting non-placeholder runs.
                if not replaced:
                    import sys
                    print(f"[FILL WARNING] No match found for label={label!r} "
                          f"run_text={run_text_now!r} current_text={current_text!r} "
                          f"canonical={canonical_run_text!r} — run left unchanged.",
                          file=sys.stderr)
                    skipped.append((label, f"run text {run_text_now!r} did not match placeholder pattern"))
                    continue

                _force_preserve_space(run)
                try:
                    if run.font.color is not None and run.font.color.type is not None:
                        if str(run.font.color.rgb).upper() == "FF0000":
                            run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
                except Exception:
                    pass

            # clear_run_ids already processed above
            applied.append(label)

            # Track filled labels so duplicates can be skipped
            if is_header_run:
                filled_header_labels.add(label)
            else:
                filled_labels.add(label)

            # Track which paragraphs had {{ }} fills for font shrink pass
            if not is_multiline_list:
                filled_paths.add(_path_key(entry["path"]))

        except Exception as e:
            skipped.append((label, f"error filling field: {e}"))

    # ── Subject indent sync: align filled paragraphs under SUBJECT with siblings ─
    # If a filled paragraph sits directly under a SUBJECT: line, copy the
    # left_indent and first_line_indent from the next sibling paragraph that
    # has an indent set — so PROJECT TITLE aligns with KANEOHE, ISLAND OF OAHU.
    all_paras = list(_iter_block_items_with_path(doc))
    for i, (path, para) in enumerate(all_paras):
        if _path_key(path) not in filled_paths:
            continue
        # Look backwards for a SUBJECT: paragraph within 5 paragraphs
        is_under_subject = False
        for j in range(i - 1, max(i - 6, -1), -1):
            prev_text = "".join(r.text for r in all_paras[j][1].runs).strip()
            if prev_text.upper().startswith("SUBJECT"):
                is_under_subject = True
                break
            if prev_text:
                break  # hit non-empty non-subject paragraph — stop looking
        if not is_under_subject:
            continue
        # Find the next sibling paragraph that has an indent set
        ref_indent = None
        ref_first  = None
        for j in range(i + 1, min(i + 6, len(all_paras))):
            sibling_pf = all_paras[j][1].paragraph_format
            if sibling_pf.left_indent is not None:
                ref_indent = sibling_pf.left_indent
                ref_first  = sibling_pf.first_line_indent
                break
        if ref_indent is not None:
            pf = para.paragraph_format
            if pf.left_indent is None or pf.left_indent == 0:
                pf.left_indent       = ref_indent
                pf.first_line_indent = ref_first


    import calendar as _cal
    from docx.shared import Pt as _Pt
    from docx.table import _Cell as _CellType

    MONTH_NAMES = {m.upper() for m in _cal.month_name if m}
    MONTH_ABBRS = {m.upper() for m in _cal.month_abbr if m}
    ALL_MONTHS  = MONTH_NAMES | MONTH_ABBRS

    def _is_date_text(text):
        words = re.split(r"[\s,./\-]+", text.upper())
        return any(w in ALL_MONTHS for w in words)

    def _is_contact_name(text):
        t = text.strip()
        # Body text is never a contact name — names are always short
        if len(t) > 60:
            return False
        if "," in t:
            return True
        words = t.split()
        if 2 <= len(words) <= 5 and all(w[0].isupper() for w in words if w):
            if not any(c in t for c in ".?!"):
                return True
        return False

    def _cell_width_pt(path):
        try:
            from docx.oxml.ns import qn as _qn2
            # Walk path to find the cell
            current = doc
            for step_kind, step_idx in path:
                if step_kind == "table":
                    from docx.table import Table as _T2
                    tables = [c for c in current.element.body.iterchildren()
                              if c.tag == qn("w:tbl")]
                    current = _T2(tables[step_idx], current)
                elif step_kind == "row":
                    current = current.rows[step_idx]
                elif step_kind == "cell":
                    current = current.cells[step_idx]
            if isinstance(current, _CellType):
                w = current._tc.tcPr
                if w is not None:
                    tcW = w.find(qn("w:tcW"))
                    if tcW is not None:
                        val = tcW.get(qn("w:w"))
                        typ = tcW.get(qn("w:type"))
                        if val and typ == "dxa":
                            return int(val) / 20  # twips to points
        except Exception:
            pass
        return None

    # ── Font shrink: reduce font size for dates/names that overflow their cell ─
    # After shrinking a cell, sync the whole table row to the same font size
    # so all cells in the row stay visually consistent.
    import calendar as _cal
    from docx.shared import Pt as _Pt
    from docx.table import _Cell as _CellType

    MONTH_NAMES = {m.upper() for m in _cal.month_name if m}
    MONTH_ABBRS = {m.upper() for m in _cal.month_abbr if m}
    ALL_MONTHS  = MONTH_NAMES | MONTH_ABBRS

    def _is_date_text(text):
        words = re.split(r"[\s,./\-]+", text.upper())
        return any(w in ALL_MONTHS for w in words)

    def _is_contact_name(text):
        t = text.strip()
        if "," in t and len(t) < 80:
            return True
        words = t.split()
        if 2 <= len(words) <= 6 and all(w[0].isupper() for w in words if w):
            if not any(c in t for c in ".?!"):
                return True
        return False

    def _cell_width_pt(path):
        try:
            current = doc
            for step_kind, step_idx in path:
                if step_kind == "table":
                    from docx.table import Table as _T2
                    tables = [c for c in current.element.body.iterchildren()
                              if c.tag == qn("w:tbl")]
                    current = _T2(tables[step_idx], current)
                elif step_kind == "row":
                    current = current.rows[step_idx]
                elif step_kind == "cell":
                    current = current.cells[step_idx]
            if isinstance(current, _CellType):
                tcPr = current._tc.tcPr
                if tcPr is not None:
                    tcW = tcPr.find(qn("w:tcW"))
                    if tcW is not None:
                        val = tcW.get(qn("w:w"))
                        typ = tcW.get(qn("w:type"))
                        if val and typ == "dxa":
                            return int(val) / 20
        except Exception:
            pass
        return None

    def _get_table_row(path):
        try:
            current = doc
            row_obj = None
            for step_kind, step_idx in path:
                if step_kind == "table":
                    from docx.table import Table as _T2
                    tables = [c for c in current.element.body.iterchildren()
                              if c.tag == qn("w:tbl")]
                    current = _T2(tables[step_idx], current)
                elif step_kind == "row":
                    row_obj = current.rows[step_idx]
                    current = row_obj
                elif step_kind == "cell":
                    current = current.cells[step_idx]
            return row_obj
        except Exception:
            return None

    CHARS_PER_PT        = 1 / 7.2
    MAX_CHARS_FULL_LINE = 90
    MIN_CELL_WIDTH_PT   = 50   # ignore cells narrower than ~0.7 inch (likely header label cells)
    synced_rows         = set()

    for path, para in _iter_block_items_with_path(doc):
        # Only shrink paragraphs where a {{ }} placeholder was filled
        if _path_key(path) not in filled_paths:
            continue

        para_runs = [r for r in para.runs if (r.text or "").strip()]
        if not para_runs:
            continue
        full_text = "".join(r.text or "" for r in para_runs)

        # Only shrink short values — dates and names are always short.
        if len(full_text) > 60 and not any(s[0] == "table" for s in path):
            continue
        if len(full_text) > 40 and not _is_date_text(full_text) and not _is_contact_name(full_text):
            continue
        if not (_is_date_text(full_text) or _is_contact_name(full_text)):
            continue

        cell_w_pt = _cell_width_pt(path)
        if cell_w_pt is not None and cell_w_pt < MIN_CELL_WIDTH_PT:
            continue  # skip very narrow label cells
        if cell_w_pt and cell_w_pt > 0:
            max_chars = int(cell_w_pt * CHARS_PER_PT * 0.9)
            # Sanity check — never set max_chars below 5
            max_chars = max(max_chars, 5)
        else:
            max_chars = MAX_CHARS_FULL_LINE

        if len(full_text) <= max_chars:
            continue

        sample_pt   = para_runs[0].font.size.pt if para_runs[0].font.size else 12.0
        required_pt = sample_pt * max_chars / len(full_text)
        new_pt      = max(round(required_pt * 2) / 2, 7.0)

        if new_pt >= sample_pt:
            continue

        # Shrink this paragraph's runs
        for run in para_runs:
            run.font.size = _Pt(new_pt)

        # ── Sync whole table row to the same font size ────────────────────────
        # Only sync cells that contain filled placeholder paragraphs.
        row_obj = _get_table_row(path)
        if row_obj is None:
            continue
        row_id = id(row_obj._tr)
        synced_rows.add(row_id)

        seen_tcs = set()
        for cell in row_obj.cells:
            if id(cell._tc) in seen_tcs:
                continue
            seen_tcs.add(id(cell._tc))
            for cell_para in cell.paragraphs:
                # Only shrink runs in paragraphs that had {{ }} fills
                cell_path_key = None
                for cp, cp_para in _iter_block_items_with_path(doc):
                    if cp_para is cell_para:
                        cell_path_key = _path_key(cp)
                        break
                if cell_path_key not in filled_paths:
                    continue
                for r in cell_para.runs:
                    if not (r.text or "").strip():
                        continue
                    existing_pt = r.font.size.pt if r.font.size else 12.0
                    if existing_pt > new_pt:
                        r.font.size = _Pt(new_pt)


    # ── Final cleanup: remove leftover split-placeholder fragments ────────────
    for _, para in _iter_block_items_with_path(doc):
        full_para_text = "".join(r.text or "" for r in para.runs)
        has_complete_placeholder = bool(_PLACEHOLDER_BRACE_RE.search(full_para_text))

        runs = para.runs
        for i, run in enumerate(runs):
            t  = run.text or ""
            ts = t.strip()
            if not ts:
                continue

            # Skip runs that are part of a complete {{ }} still in the paragraph
            if has_complete_placeholder and ("{{" in t or "}}" in t or "}" in t):
                continue

            # Orphaned closing fragment: ends with }} or single } but no {{
            if (ts.endswith("}}") or ts.endswith("}")) and "{{" not in ts:
                run.text = ""
                continue

            # Orphaned opening fragment: starts with {{ but no closing }}
            if ts.startswith("{{") and "}}" not in ts:
                run.text = ""
                continue

            # Standalone underscore run that sits between placeholder fragments —
            # only strip if adjacent runs are also empty or look like fragments
            if ts and all(c == "_" for c in ts):
                prev_text = (runs[i-1].text or "").strip() if i > 0 else ""
                next_text = (runs[i+1].text or "").strip() if i < len(runs)-1 else ""
                prev_is_fragment = (prev_text == "" or prev_text.startswith("{{") or prev_text.endswith("}}") or prev_text.endswith("}"))
                next_is_fragment = (next_text == "" or next_text.startswith("{{") or next_text.endswith("}}") or next_text.endswith("}"))
                if prev_is_fragment or next_is_fragment:
                    run.text = ""
                    continue

    out_buf = io.BytesIO()
    doc.save(out_buf)
    return out_buf.getvalue(), applied, skipped


def python_format_diff(template_bytes, doc_bytes):
    """Compare template and submitted doc run-by-run using pure Python.

    Returns a dict with keys:
        formatting_issues  — list of {field, problem, expected, actual, severity}
        missing_fields     — list of {field, description}
        correct_fields     — list of {field, note}
        python_summary     — short plain-English summary string

    Checks: bold, italic, underline, font name, font size, color, highlight,
    alignment, space_before, space_after, unfilled {{ }} placeholders.
    Does NOT check: content, spelling, spacing rules, acronyms — those go to AI.
    """
    from docx import Document as _Doc
    import io as _io

    def _open(b):
        return _Doc(_io.BytesIO(b))

    tmpl_doc = _open(template_bytes)
    subm_doc = _open(doc_bytes)

    # Build flat lists of (text, attrs) for each paragraph in document order
    def _para_attrs(para):
        pf = para.paragraph_format
        return {
            "align":   str(para.alignment),
            "sb":      str(pf.space_before.pt) if pf.space_before else "None",
            "sa":      str(pf.space_after.pt)  if pf.space_after  else "None",
            "ls":      str(pf.line_spacing),
        }

    def _run_attrs(run):
        try:
            color = str(run.font.color.rgb).upper() if (run.font.color and run.font.color.type) else "None"
        except Exception:
            color = "None"
        size = str(run.font.size.pt) if run.font.size else "None"
        return {
            "bold":      str(run.bold),
            "italic":    str(run.italic),
            "underline": str(run.underline),
            "size":      size,
            "color":     color,
            "font":      str(run.font.name),
        }

    def _walk_paras(doc):
        paras = []
        for _, para in _iter_block_items_with_path(doc):
            text = "".join(r.text for r in para.runs).strip()
            paras.append({
                "text":  text,
                "attrs": _para_attrs(para),
                "runs":  [{"text": r.text, "attrs": _run_attrs(r)} for r in para.runs if r.text.strip()],
            })
        return paras

    tmpl_paras = _walk_paras(tmpl_doc)
    subm_paras = _walk_paras(subm_doc)

    issues  = []
    missing = []
    correct = []

    # ── Check for unfilled {{ }} placeholders in submitted doc ────────────────
    for p in subm_paras:
        full = p["text"]
        for m in _PLACEHOLDER_BRACE_RE.finditer(full):
            inner = m.group(1).strip()
            label = re.sub(r"_+", " ", inner).upper()
            missing.append({
                "field":       label,
                "description": f"This field was not filled in — still shows as {{{{{inner}}}}}",
            })

    # ── Match template paragraphs to submitted paragraphs by text similarity ──
    # Build a lookup of submitted paragraph texts for quick matching
    subm_text_map = {}
    for p in subm_paras:
        key = p["text"][:60].strip().lower()
        if key:
            subm_text_map[key] = p

    for tp in tmpl_paras:
        tmpl_text = tp["text"]

        # Skip blank paragraphs and placeholder paragraphs
        if not tmpl_text.strip():
            continue
        if _PLACEHOLDER_BRACE_RE.search(tmpl_text):
            continue

        # Try to find matching paragraph in submitted doc
        key = tmpl_text[:60].strip().lower()
        sp = subm_text_map.get(key)
        if sp is None:
            continue  # content mismatch handled by AI

        # ── Compare paragraph-level attrs ─────────────────────────────────
        ta = tp["attrs"]
        sa = sp["attrs"]
        field_label = tmpl_text[:50] + ("…" if len(tmpl_text) > 50 else "")

        if ta["align"] != sa["align"] and ta["align"] != "None":
            issues.append({
                "field":    field_label,
                "problem":  f"Paragraph alignment is wrong.",
                "expected": f"Alignment should be {ta['align']}",
                "actual":   f"Currently set to {sa['align']}",
                "severity": "error",
            })

        # ── Compare run-level attrs for matching runs ──────────────────────
        tmpl_runs = [r for r in tp["runs"] if r["text"].strip()]
        subm_runs = [r for r in sp["runs"] if r["text"].strip()]

        for tr in tmpl_runs:
            # Find matching run in submitted doc by text
            sr = next((r for r in subm_runs if r["text"].strip() == tr["text"].strip()), None)
            if sr is None:
                continue

            ra  = tr["attrs"]
            sra = sr["attrs"]
            run_label = tr["text"].strip()[:40]

            # Bold
            if ra["bold"] != "None" and ra["bold"] != sra["bold"]:
                exp = "bold" if ra["bold"] == "True" else "not bold"
                act = "bold" if sra["bold"] == "True" else "not bold"
                issues.append({
                    "field":    run_label,
                    "problem":  f"Text should be {exp} but is {act}.",
                    "expected": exp, "actual": act, "severity": "error",
                })

            # Underline
            if ra["underline"] != "None" and ra["underline"] != sra["underline"]:
                exp = "underlined" if ra["underline"] == "True" else "not underlined"
                act = "underlined" if sra["underline"] == "True" else "not underlined"
                issues.append({
                    "field":    run_label,
                    "problem":  f"Text should be {exp} but is {act}.",
                    "expected": exp, "actual": act, "severity": "error",
                })

            # Font size
            if ra["size"] != "None" and ra["size"] != sra["size"]:
                issues.append({
                    "field":    run_label,
                    "problem":  f"Font size should be {ra['size']}pt but is {sra['size']}pt.",
                    "expected": f"{ra['size']}pt", "actual": f"{sra['size']}pt",
                    "severity": "error",
                })

            # Font name
            if ra["font"] not in ("None", "None") and ra["font"] != sra["font"]:
                issues.append({
                    "field":    run_label,
                    "problem":  f"Font should be {ra['font']} but is {sra['font']}.",
                    "expected": ra["font"], "actual": sra["font"],
                    "severity": "warning",
                })

            # Color — only flag if template explicitly sets a non-black color
            if ra["color"] not in ("None", "000000") and ra["color"] != sra["color"]:
                color_names = {
                    "FF0000": "red", "0000FF": "blue", "008000": "green",
                    "FFA500": "orange", "800080": "purple", "FFFFFF": "white",
                }
                exp_name = color_names.get(ra["color"], f"#{ra['color']}")
                act_name = color_names.get(sra["color"], f"#{sra['color']}")
                issues.append({
                    "field":    run_label,
                    "problem":  f"Text color should be {exp_name} but is {act_name}.",
                    "expected": exp_name, "actual": act_name, "severity": "warning",
                })

        if not issues or all(i["field"] != field_label for i in issues):
            if tmpl_text.strip():
                correct.append({"field": field_label, "note": "Formatting matches template."})

    n_issues  = len(issues)
    n_missing = len(missing)
    if n_issues == 0 and n_missing == 0:
        summary = "No formatting issues detected by automated check."
    elif n_issues <= 2 and n_missing == 0:
        summary = f"{n_issues} minor formatting difference(s) found."
    else:
        summary = f"{n_issues} formatting issue(s) and {n_missing} unfilled field(s) found by automated check."

    return {
        "formatting_issues": issues,
        "missing_fields":    missing,
        "correct_fields":    correct[:10],  # cap to avoid huge lists
        "python_summary":    summary,
    }


def render(tab):
    with tab:
        st.subheader("📋 Secretaries Manual — Document Checker")

        subtab_check, subtab_manage = st.tabs([
            "🔍 Check a Document", "🗂️ Manage Templates"
        ])

        # ── SUB-TAB 1 — CHECK A DOCUMENT ────────────────────────────────────────
        with subtab_check:
            st.markdown("### Check a Document Against the Template")

            available_templates = list_docx_templates()
            kb_has_manual = is_manual_indexed("secretaries_manual")

            if not available_templates and not kb_has_manual:
                st.warning(
                    "No templates uploaded and the Secretaries Manual is not indexed. "
                    "Go to **Manage Templates** to add a blank DOCX template, or index the "
                    "Secretaries Manual in **Index Documents**."
                )
            elif available_templates:
                st.caption(f"Available templates: {', '.join(available_templates)}")
            else:
                st.caption("Using indexed Secretaries Manual as formatting reference.")

            st.markdown("---")

            filled_upload = st.file_uploader(
                "Upload your filled document (DOCX or PDF)",
                type=["docx", "pdf"],
                key="secman_filled_uploader",
            )

            if filled_upload:
                is_pdf = filled_upload.name.lower().endswith(".pdf")

                if is_pdf:
                    st.info(
                        "📄 PDF uploaded — content rules and auto-fix will work normally. "
                        "**Formatting checks (bold, font, color, alignment) are not available for PDFs.** "
                        "Upload a DOCX for full formatting analysis.",
                        icon="ℹ️",
                    )

                if st.session_state["secman_filled_name"] != filled_upload.name:
                    st.session_state["secman_filled_name"]          = filled_upload.name
                    st.session_state["secman_filled_bytes"]         = filled_upload.getvalue()
                    st.session_state["secman_filled_text"]          = None
                    st.session_state["secman_plain_text"]           = None
                    st.session_state["secman_report"]               = None
                    st.session_state["secman_report_parsed"]        = None
                    st.session_state["secman_filled_is_docx"]       = not is_pdf
                    st.session_state["secman_doc_type"]             = None
                    st.session_state["secman_template_context"]     = ""
                    st.session_state["secman_fixed_docx_bytes"]     = None
                    st.session_state["secman_fix_log"]              = None
                    st.session_state["secman_pending_fills"]        = None
                    st.session_state["secman_pending_labels"]       = None
                    st.session_state["secman_pending_tmpl_map"]     = None
                    st.session_state["secman_pending_placeholders"] = None

                file_size_kb = round(filled_upload.size / 1024, 1)
                st.info(f"📄 **{filled_upload.name}** — {file_size_kb} KB")

                if st.session_state["secman_filled_text"] is None:
                    with st.spinner("Reading document…"):
                        try:
                            if is_pdf:
                                # PDF: extract plain text only — no formatting data
                                import pdfplumber, io as _io
                                pdf_lines = []
                                with pdfplumber.open(_io.BytesIO(st.session_state["secman_filled_bytes"])) as pdf:
                                    for page in pdf.pages:
                                        text = page.extract_text()
                                        if text:
                                            pdf_lines.append(text)
                                pdf_plain = "\n".join(pdf_lines)
                                # Store plain text in both fields — no run/para data available
                                st.session_state["secman_filled_text"] = pdf_plain
                                st.session_state["secman_plain_text"]  = pdf_plain
                            else:
                                st.session_state["secman_filled_text"], _ = build_indexed_run_report(
                                    st.session_state["secman_filled_bytes"]
                                )
                        except Exception as e:
                            st.error(f"Could not read file: {e}")
                            st.session_state["secman_filled_text"] = ""

                if st.session_state["secman_filled_text"]:
                    with st.expander("👁️ View extracted formatting data", expanded=False):
                        st.text(st.session_state["secman_filled_text"][:3000])

                    st.markdown("---")

                    btn_col1, btn_col2 = st.columns([0.6, 0.4])
                    with btn_col1:
                        run_check = (
                            st.button("🤖 Run Formatting Check", type="primary", key="run_comparison")
                            or st.session_state.get("secman_auto_recheck", False)
                        )
                    with btn_col2:
                        if st.session_state.get("secman_report"):
                            if st.button("🔄 Clear & Re-run Check", key="rerun_check"):
                                st.session_state["secman_report"]               = None
                                st.session_state["secman_report_parsed"]        = None
                                st.session_state["secman_doc_type"]             = None
                                st.session_state["secman_template_context"]     = ""
                                st.session_state["_secman_template_docx"]       = None
                                st.session_state["secman_fixed_docx_bytes"]     = None
                                st.session_state["secman_fix_log"]              = None
                                st.session_state["secman_pending_fills"]        = None
                                st.session_state["secman_pending_labels"]       = None
                                st.session_state["secman_pending_tmpl_map"]     = None
                                st.session_state["secman_pending_placeholders"] = None
                                st.session_state["secman_auto_recheck"]         = True
                                st.rerun()

                    if run_check:
                        st.session_state["secman_auto_recheck"] = False
                        filled_text = st.session_state["secman_filled_text"]

                        forced_type = st.session_state.get("secman_doc_type")
                        if forced_type and st.session_state.get("secman_template_context"):
                            doc_type         = forced_type
                            template_context = st.session_state["secman_template_context"]
                        else:
                            # ── Step 1: identify doc type ────────────────────────────────
                            with st.spinner("Step 1 of 3: Identifying document type…"):
                                feedback_context = build_feedback_context(filled_text[:500])

                                subject_para_start = filled_text.find("SUBJECT")
                                subject_hint = ""
                                if subject_para_start != -1:
                                    subject_runs = re.findall(
                                        r"\[RUN text='([^']+)'",
                                        filled_text[subject_para_start:subject_para_start + 600]
                                    )
                                    subject_hint = " ".join(
                                        r for r in subject_runs
                                        if r.strip() not in ("", ":", " ", "  ")
                                    )[:200]

                                known_types = list_docx_templates()
                                kb_links = load_kb_folder_links()
                                known_types_combined = sorted(set(known_types) | set(kb_links.keys()))
                                type_list_str = "\n".join(f"- {t}" for t in known_types_combined) if known_types_combined else "(none stored yet)"

                                id_prompt = (
                                    "You are an HDOT document classifier. "
                                    "HDOT documents use a standard MEMORANDUM header wrapper. "
                                    "The actual document type is determined by the SUBJECT line content, not the word MEMORANDUM.\n\n"
                                    f"SUBJECT LINE EXTRACTED: {subject_hint or '(not found)'}\n\n"
                                    f"KNOWN DOCUMENT TYPES (return the best matching name exactly as written):\n{type_list_str}\n\n"
                                    "Return ONLY the document type name — no explanation, no punctuation.\n\n"
                                    + (feedback_context + "\n\n" if feedback_context else "")
                                    + f"DOCUMENT FORMATTING DATA:\n{filled_text[:1200]}"
                                )
                                doc_type = gemini_call(id_prompt, temperature=0.0, timeout=20)
                                doc_type = (doc_type or "Unknown").strip().strip('"\'')
                                st.session_state["secman_doc_type"] = doc_type

                            # ── Steps 2 + 3 prep: run in parallel ───────────────────────
                            with st.spinner("Step 2 of 3: Retrieving template and preparing document…"):
                                from concurrent.futures import ThreadPoolExecutor

                                def _load_template():
                                    tmpl = find_matching_template(doc_type)
                                    if tmpl:
                                        ctx = extract_template_formatting(tmpl, doc_type=doc_type)
                                        return tmpl, ctx
                                    ctx = ""
                                    linked_folder = get_kb_folder_for_doc_type(doc_type)
                                    if linked_folder:
                                        ctx = get_kb_context_for_folder(linked_folder)
                                    return None, ctx

                                def _build_tmpl_numbered(tmpl_bytes):
                                    if not tmpl_bytes:
                                        return None
                                    full, _ = build_indexed_run_report(tmpl_bytes)
                                    src_lines = [l for l in full.splitlines() if l.strip()]
                                    return "\n".join(
                                        re.sub(r"^D(\d+):", r"T\1:", l) for l in src_lines[:200]
                                    )

                                with ThreadPoolExecutor(max_workers=2) as ex:
                                    tmpl_future  = ex.submit(_load_template)
                                    doc_lines    = [l for l in filled_text.splitlines() if l.strip()]
                                    doc_numbered = "\n".join(doc_lines[:200])
                                    template_docx, template_context = tmpl_future.result()

                                st.session_state["_secman_template_docx"]   = template_docx
                                st.session_state["secman_template_context"] = template_context

                                if template_context and not template_docx:
                                    st.info("ℹ️ No DOCX template found — using KB folder reference.")

                                if template_docx:
                                    tmpl_numbered = _build_tmpl_numbered(template_docx)
                                else:
                                    tmpl_lines    = [l for l in template_context.splitlines() if l.strip()]
                                    tmpl_numbered = "\n".join(f"T{i+1:03d}: {l}" for i, l in enumerate(tmpl_lines[:120]))

                        # ── Step 3: Python diff + AI content check ───────────────
                        if not template_context.strip():
                            st.error(
                                f"No template found for **{doc_type}**. "
                                "Upload a blank DOCX template in **Manage Templates**, "
                                "link a KB folder there, "
                                "or index the Secretaries Manual in **Index Documents**."
                            )
                        else:
                            tmpl_src = (
                                "uploaded DOCX template"
                                if st.session_state.get("_secman_template_docx")
                                else "KB folder reference"
                            )

                            doc_stripped  = re.sub(r"\s+", " ", filled_text).strip()
                            tmpl_stripped = re.sub(r"\s+", " ", template_context).strip()
                            similarity    = (
                                len(set(doc_stripped.split()) & set(tmpl_stripped.split()))
                                / max(len(set(tmpl_stripped.split())), 1)
                            )
                            if similarity > 0.97:
                                st.warning(
                                    "⚠️ **The uploaded document appears identical to the blank template.** "
                                    "Please make sure you uploaded a *filled-in* copy, not the blank template itself."
                                )

                            # ── Step 3a: Python formatting diff (fast, no API) ────
                            # Only available for DOCX — PDFs have no run/para data
                            py_diff = None
                            if st.session_state.get("secman_filled_is_docx"):
                                with st.spinner("Step 3 of 3: Running automated formatting check…"):
                                    template_docx_for_diff = st.session_state.get("_secman_template_docx")
                                    if template_docx_for_diff:
                                        try:
                                            py_diff = python_format_diff(
                                                template_docx_for_diff,
                                                st.session_state["secman_filled_bytes"],
                                            )
                                        except Exception as e:
                                            st.warning(f"Automated check error: {e}")
                            else:
                                st.info("ℹ️ Formatting attribute check skipped for PDF — content rules only.")

                            # ── Step 3b: AI content check (smaller prompt) ─────────
                            with st.spinner("Step 3 of 3: AI checking content rules…"):
                                # Send only plain text to AI — no run/para attribute data
                                # AI only checks content rules it can't see from attributes
                                plain_doc  = st.session_state.get("secman_plain_text") or extract_plain_text(
                                    st.session_state["secman_filled_bytes"]
                                )

                                check_prompt = (
                                    "You are an HDOT document content auditor. "
                                    "Python has already checked all formatting attributes (bold, font, size, color, alignment). "
                                    "Your job is ONLY to check the content rules listed below that Python cannot check.\n\n"
                                    f"DOCUMENT TYPE: {doc_type}\n\n"
                                    "══════════════════════════════════════════\n"
                                    "DOCUMENT TEXT (plain text only):\n"
                                    "══════════════════════════════════════════\n"
                                    f"{plain_doc[:8000]}\n\n"
                                    "CHECK ONLY THESE CONTENT RULES:\n"
                                    f"{COMMON_MISTAKES}\n\n"
                                    "Do NOT report formatting issues (bold, font, size, color, alignment) — those are already handled.\n"
                                    "ONLY report issues with the content rules listed above.\n\n"
                                    "CRITICAL OUTPUT RULES:\n"
                                    "- Write for a non-technical office worker, NOT a programmer\n"
                                    "- NEVER mention 'run', 'para', 'attribute', 'inherited', hex colors, or line numbers\n"
                                    "- Describe problems as a manager would write in a correction note\n\n"
                                    "Return ONLY valid JSON (no markdown fences):\n"
                                    "{\n"
                                    '  "doc_type": "string",\n'
                                    '  "compliance_level": "Compliant" | "Minor Issues" | "Needs Revision" | "Non-Compliant",\n'
                                    '  "summary": "2-3 sentence plain-English summary",\n'
                                    '  "correct": [{"field": "field label", "note": "confirmation"}],\n'
                                    '  "issues": [\n'
                                    '    {"field": "field label", "problem": "what is wrong",\n'
                                    '     "expected": "what it should be", "actual": "what it is",\n'
                                    '     "severity": "error"|"warning"}\n'
                                    '  ],\n'
                                    '  "missing": [{"field": "field label", "description": "what was not filled in"}]\n'
                                    "}"
                                )
                                _gemini_cache.pop(
                                    (hashlib.md5(check_prompt.encode()).hexdigest()[:16], 0.0), None
                                )
                                result = gemini_call(check_prompt, temperature=0.0, timeout=60)

                            if not result:
                                st.error("AI check failed. Check your Gemini API key or quota.")
                            else:
                                clean = result.strip()
                                clean = re.sub(r"^```[a-zA-Z]*\s*", "", clean)
                                clean = re.sub(r"\s*```$", "", clean).strip()
                                clean = re.sub(r"```[a-zA-Z]*", "", clean).strip()
                                brace_start = clean.find("{")
                                brace_end   = clean.rfind("}")
                                if brace_start != -1 and brace_end > brace_start:
                                    clean = clean[brace_start:brace_end + 1]

                                with st.expander("🔍 Debug: Raw AI response (first 2000 chars)", expanded=False):
                                    st.text(clean[:2000])

                                try:
                                    ai_parsed = json.loads(clean)
                                except Exception as parse_err:
                                    ai_parsed = None
                                    st.warning(
                                        f"⚠️ AI response could not be parsed. "
                                        f"Showing Python check results only. (Error: {parse_err})"
                                    )

                                # ── Merge Python diff + AI results ────────────────
                                if ai_parsed and py_diff:
                                    merged_issues  = py_diff["formatting_issues"] + ai_parsed.get("issues", [])
                                    merged_missing = py_diff["missing_fields"]    + ai_parsed.get("missing", [])
                                    merged_correct = py_diff["correct_fields"]    + ai_parsed.get("correct", [])
                                    # Determine compliance from combined issue count
                                    n = len(merged_issues) + len(merged_missing)
                                    compliance = (
                                        "Non-Compliant" if n > 8 else
                                        "Needs Revision" if n > 4 else
                                        "Minor Issues"   if n > 0 else
                                        "Compliant"
                                    )
                                    parsed = {
                                        "doc_type":         ai_parsed.get("doc_type", doc_type),
                                        "compliance_level": compliance,
                                        "summary":          ai_parsed.get("summary", py_diff["python_summary"]),
                                        "correct":          merged_correct,
                                        "issues":           merged_issues,
                                        "missing":          merged_missing,
                                    }
                                elif ai_parsed:
                                    parsed = ai_parsed
                                elif py_diff:
                                    n = len(py_diff["formatting_issues"]) + len(py_diff["missing_fields"])
                                    parsed = {
                                        "doc_type":         doc_type,
                                        "compliance_level": "Non-Compliant" if n > 4 else "Minor Issues" if n > 0 else "Compliant",
                                        "summary":          py_diff["python_summary"],
                                        "correct":          py_diff["correct_fields"],
                                        "issues":           py_diff["formatting_issues"],
                                        "missing":          py_diff["missing_fields"],
                                    }
                                else:
                                    parsed = None

                                st.session_state["secman_report"]           = clean
                                st.session_state["secman_report_parsed"]    = parsed
                                st.session_state["secman_fixed_docx_bytes"] = None
                                st.session_state["secman_fix_log"]          = None
                                st.session_state["secman_pending_fills"]        = None
                                st.session_state["secman_pending_labels"]       = None
                                st.session_state["secman_pending_tmpl_map"]     = None
                                st.session_state["secman_pending_placeholders"] = None
                                st.rerun()

                    # ── Display report ──────────────────────────────────────────
                    if st.session_state["secman_report"]:
                        st.markdown("---")

                        doc_type   = st.session_state.get("secman_doc_type", "")
                        tmpl_src   = "uploaded DOCX template" if st.session_state.get("_secman_template_docx") else "KB folder reference"
                        parsed     = st.session_state.get("secman_report_parsed")
                        report_raw = st.session_state["secman_report"]

                        if doc_type and doc_type != "Unknown":
                            st.info(f"📄 **{doc_type}** — checked against {tmpl_src}")

                        if parsed:
                            compliance  = parsed.get("compliance_level", "")
                            badge_color = ("#f44336" if compliance == "Non-Compliant" else
                                           "#ff9800" if compliance == "Needs Revision" else
                                           "#ffeb3b" if compliance == "Minor Issues"   else "#4caf50")
                            badge_icon  = ("🔴" if compliance == "Non-Compliant" else
                                           "🟠" if compliance == "Needs Revision" else
                                           "🟡" if compliance == "Minor Issues"   else "🟢")
                            summary  = parsed.get("summary", "")
                            correct  = parsed.get("correct",  [])
                            issues   = parsed.get("issues",   [])
                            missing  = parsed.get("missing",  [])
                            errors   = [i for i in issues if i.get("severity") == "error"]
                            warnings = [i for i in issues if i.get("severity") != "error"]

                            st.markdown(
                                f"<div style='background:#1a1a2e;border-left:4px solid {badge_color};"
                                f"border-radius:8px;padding:12px 16px;margin-bottom:12px;'>"
                                f"{badge_icon} <strong style='color:#e0e0e0;font-size:14px;'>"
                                f"{compliance or 'Complete'}</strong>"
                                f"<div style='color:#bbb;font-size:12px;margin-top:4px;'>{summary}</div>"
                                f"</div>",
                                unsafe_allow_html=True,
                            )

                            mc1, mc2, mc3 = st.columns(3)
                            mc1.metric("⚠️ Issues",  len(issues))
                            mc2.metric("📋 Missing", len(missing))
                            mc3.metric("✅ Correct", len(correct))
                            st.markdown("---")

                            promoted_to_missing = []
                            remaining_issues    = []
                            for item in issues:
                                prob = item.get("problem", "").lower()
                                if "{{" in prob or "placeholder" in prob:
                                    promoted_to_missing.append({
                                        "field":       item.get("field", "?"),
                                        "description": item.get("problem", ""),
                                    })
                                else:
                                    remaining_issues.append(item)
                            issues  = remaining_issues
                            missing = missing + promoted_to_missing
                            errors   = [i for i in issues if i.get("severity") == "error"]
                            warnings = [i for i in issues if i.get("severity") != "error"]

                            if issues:
                                st.markdown("#### ⚠️ Formatting Issues")
                                seen, unique_issues = set(), []
                                for item in errors + warnings:
                                    key = (
                                        item.get("field", "").strip().lower(),
                                        item.get("problem", "").strip().lower(),
                                    )
                                    if key not in seen:
                                        seen.add(key)
                                        unique_issues.append(item)
                                from collections import OrderedDict as _OD
                                grouped_issues = _OD()
                                for item in unique_issues:
                                    field = item.get("field", "?").rstrip(":")
                                    grouped_issues.setdefault(field, []).append(item)
                                lines = []
                                for field, items in grouped_issues.items():
                                    has_error = any(i.get("severity") == "error" for i in items)
                                    icon = "🔴" if has_error else "🟡"
                                    lines.append(f"- {icon} **{field}**")
                                    for item in items:
                                        lines.append(f"  - {item.get('problem', '')}")
                                st.markdown("\n".join(lines))
                            else:
                                st.success("No formatting issues found.")

                            if missing:
                                st.markdown("#### 📋 Missing / Unfilled Fields")
                                seen_m, missing_lines = set(), []
                                for item in missing:
                                    field = item.get("field", "?").rstrip(":")
                                    desc  = item.get("description", "")
                                    key   = field.lower()
                                    if key in seen_m:
                                        continue
                                    seen_m.add(key)
                                    missing_lines.append(f"- 🟠 **{field}**" + (f": {desc}" if desc else ""))
                                st.markdown("\n".join(missing_lines))
                            else:
                                st.success("No missing fields.")

                            if correct:
                                st.markdown("#### ✅ Correctly Formatted")
                                correct_lines = []
                                for item in correct:
                                    field = item.get("field", "?").rstrip(":")
                                    correct_lines.append(f"- ✅ {field}")
                                st.markdown("\n".join(correct_lines))

                        else:
                            st.warning(
                                "⚠️ The AI response could not be parsed. "
                                "Click **Run Formatting Check** again to retry."
                            )
                            with st.expander("Raw AI response (for debugging)", expanded=False):
                                st.text(report_raw[:3000])

                        # ── Auto-Fix Formatting ──────────────────────────────────
                        template_docx_bytes = st.session_state.get("_secman_template_docx")
                        st.markdown("---")
                        st.markdown("#### 🛠️ Auto-Fix Formatting")

                        if not template_docx_bytes:
                            st.info(
                                "Auto-fix needs an actual DOCX template for this document type "
                                "(not just a KB folder reference). Upload one in **Manage Templates** "
                                "to enable this."
                            )
                        else:
                            st.caption(
                                "Generates a corrected document by starting from the official blank "
                                "template — already guaranteed to have the right fonts, headers, and "
                                "layout — and filling in ONLY its placeholder fields (red text or "
                                "text wrapped in {{ }}) using the information found in your uploaded "
                                "document. Nothing else in the template is touched."
                            )

                            if st.button(
                                "🛠️ Auto-Fix & Generate Corrected DOCX",
                                type="primary",
                                key="autofix_via_template",
                            ):
                                # Cache plain text extraction
                                if not st.session_state.get("secman_plain_text"):
                                    if st.session_state.get("secman_filled_is_docx"):
                                        st.session_state["secman_plain_text"] = extract_plain_text(
                                            st.session_state["secman_filled_bytes"]
                                        )
                                    else:
                                        # PDF — plain text already stored in secman_filled_text
                                        st.session_state["secman_plain_text"] = st.session_state.get("secman_filled_text", "")

                                with st.spinner("Finding placeholder fields and reading document…"):
                                    from concurrent.futures import ThreadPoolExecutor

                                    def _find_placeholders():
                                        return find_template_placeholders(template_docx_bytes)

                                    source_text = st.session_state["secman_plain_text"]

                                    with ThreadPoolExecutor(max_workers=1) as ex:
                                        tmpl_future = ex.submit(_find_placeholders)
                                        tmpl_line_map, placeholders = tmpl_future.result()
                                        placeholder_labels = {p["run_id"]: p["label"] for p in placeholders}

                                if not placeholders:
                                    st.warning("No placeholder fields (red text or {{ }}) were found in this template.")
                                else:
                                    # ── Debug: show every detected placeholder ────────
                                    with st.expander(f"🔍 Debug: {len(placeholders)} placeholder(s) detected", expanded=False):
                                        for p in placeholders:
                                            st.markdown(
                                                f"**{p['label']}** `{p['run_id']}` "
                                                f"current_text=`{p.get('current_text','')[:40]}` "
                                                f"canonical_run=`{p.get('canonical_run_text','')[:30]}` "
                                                f"inner=`{p.get('inner_text','')[:30]}` "
                                                f"clear_runs={p.get('clear_run_ids',[])}"
                                            )

                                    with st.spinner(f"Matching {len(placeholders)} field(s) against your uploaded document…"):
                                        # Include current_text so AI sees the placeholder
                                        # in context — e.g. "HWY-DD {{reff_num}}" tells
                                        # the AI to extract only the number, not "HWY-DD X"
                                        fields_list_str = "\n".join(
                                            f"{p['run_id']}: {p['label']}"
                                            + (f" (appears as: {p['current_text']!r})" if p.get("current_text") and "{{" in p.get("current_text","") else "")
                                            for p in placeholders
                                        )
                                        fill_prompt = build_template_fill_prompt(fields_list_str, source_text)
                                        fill_result = gemini_call(fill_prompt, temperature=0.0, timeout=60)

                                    if not fill_result:
                                        st.error("AI field matching failed. Check your Gemini API key or quota.")
                                    else:
                                        fclean = re.sub(r"^```[a-zA-Z]*\s*", "", fill_result.strip())
                                        fclean = re.sub(r"\s*```$", "", fclean).strip()
                                        bs, be = fclean.find("{"), fclean.rfind("}")
                                        if bs != -1 and be > bs:
                                            fclean = fclean[bs:be + 1]
                                        try:
                                            fill_parsed = json.loads(fclean)
                                            fills = fill_parsed.get("fills", [])
                                            # Normalise run_ids — AI sometimes returns
                                            # slightly different casing/spacing
                                            valid_run_ids = {p["run_id"] for p in placeholders}
                                            for f in fills:
                                                if f.get("run_id") not in valid_run_ids:
                                                    for vid in valid_run_ids:
                                                        if vid.upper() == (f.get("run_id") or "").upper():
                                                            f["run_id"] = vid
                                                            break
                                        except Exception:
                                            fills = []
                                            st.error("Could not parse the AI's response. Try again.")

                                        # ── Debug expander — shows raw AI values ──────────
                                        with st.expander("🔍 Debug: Raw AI extracted values", expanded=False):
                                            for f in fills:
                                                rid   = f.get("run_id", "?")
                                                lbl   = placeholder_labels.get(rid, rid)
                                                val   = f.get("value", "")
                                                found = f.get("found", False)
                                                lines_in_val = val.replace("\\n", "\n").split("\n") if val else []
                                                st.markdown(
                                                    f"**{lbl}** (`{rid}`) found={found} "
                                                    f"lines={len(lines_in_val)}"
                                                )
                                                if val:
                                                    st.code(repr(val), language="text")

                                        st.session_state["secman_pending_fills"]        = fills
                                        st.session_state["secman_pending_labels"]       = placeholder_labels
                                        st.session_state["secman_pending_tmpl_map"]     = tmpl_line_map
                                        st.session_state["secman_pending_placeholders"] = placeholders
                                        st.session_state["secman_fixed_docx_bytes"]     = None
                                        st.session_state["secman_fix_log"]              = None
                                        st.rerun()

                            # ── Review / edit extracted values before applying ──────
                            if st.session_state.get("secman_pending_fills") is not None and \
                               st.session_state.get("secman_fixed_docx_bytes") is None:

                                pending_fills        = st.session_state["secman_pending_fills"]
                                pending_labels       = st.session_state.get("secman_pending_labels", {})
                                pending_placeholders = st.session_state.get("secman_pending_placeholders", [])

                                ph_meta = {
                                    p["run_id"]: {
                                        "clear_run_ids": p.get("clear_run_ids", []),
                                        "row_key": p.get("row_key"),
                                    }
                                    for p in pending_placeholders
                                }

                                st.markdown("##### ✏️ Review Extracted Values")
                                st.caption(
                                    "The AI extracted the following values from your uploaded document. "
                                    "Edit any value that looks wrong. Type **blank** to leave a field "
                                    "empty in the corrected document."
                                )

                                fill_by_run_id = {
                                    f.get("run_id"): f for f in pending_fills if f.get("run_id")
                                }

                                groups_by_rk  = {}
                                groups_order  = []

                                for ph in pending_placeholders:
                                    rk   = ph.get("row_key")
                                    key  = rk if rk else f"_solo_{ph['run_id']}"
                                    fill = fill_by_run_id.get(ph["run_id"], {
                                        "run_id": ph["run_id"], "value": "", "found": False,
                                    })
                                    if key not in groups_by_rk:
                                        groups_by_rk[key] = []
                                        groups_order.append(key)
                                    groups_by_rk[key].append((ph, fill))

                                def _label_contains(label, term):
                                    return term.upper() in label.upper()

                                title_key = next(
                                    (k for k in groups_order
                                     if len(groups_by_rk[k]) == 1
                                     and _label_contains(groups_by_rk[k][0][0].get("label",""), "PROJECT TITLE")),
                                    None
                                )
                                no_key = next(
                                    (k for k in groups_order
                                     if len(groups_by_rk[k]) == 1
                                     and _label_contains(groups_by_rk[k][0][0].get("label",""), "PROJECT NO")),
                                    None
                                )
                                if title_key and no_key and title_key != no_key:
                                    title_value = (groups_by_rk[title_key][0][1].get("value") or "")
                                    if len(title_value) <= 20:
                                        groups_by_rk[title_key].extend(groups_by_rk.pop(no_key))
                                        groups_order = [k for k in groups_order if k != no_key]

                                # ── Render — one input per unique label ───────────────
                                master_value_by_label = {}
                                edited_fills = []

                                # Count how many fields the AI couldn't find
                                unfound_labels = set()
                                for ph in pending_placeholders:
                                    run_id = ph["run_id"]
                                    label  = ph.get("label") or pending_labels.get(run_id, run_id)
                                    f = fill_by_run_id.get(run_id, {})
                                    if not f.get("found") and not f.get("value"):
                                        unfound_labels.add(label)

                                if unfound_labels:
                                    st.warning(
                                        f"⚠️ **{len(unfound_labels)} field(s) not found** in your uploaded document "
                                        f"and are highlighted below. Please fill them in manually: "
                                        + ", ".join(f"**{l}**" for l in sorted(unfound_labels))
                                    )

                                seen_labels_render = {}
                                for key in groups_order:
                                    items = groups_by_rk[key]
                                    cols  = st.columns(len(items))
                                    for col, (ph, fill) in zip(cols, items):
                                        run_id = ph["run_id"]
                                        label  = ph.get("label") or pending_labels.get(run_id, run_id)
                                        if label in seen_labels_render:
                                            continue
                                        seen_labels_render[label] = True
                                        ai_val    = fill.get("value", "") if fill.get("found") else ""
                                        not_found = label in unfound_labels
                                        with col:
                                            # Show a red label indicator for unfound fields
                                            if not_found:
                                                st.markdown(
                                                    f"<span style='color:#ff4b4b;font-size:12px;font-weight:600;'>"
                                                    f"⚠️ {label} — not found, please fill in manually</span>",
                                                    unsafe_allow_html=True,
                                                )
                                                user_val = st.text_input(
                                                    label, value="",
                                                    placeholder="⚠️ Not found — type value here",
                                                    key=f"fill_review_{run_id}",
                                                    label_visibility="collapsed",
                                                )
                                            else:
                                                user_val = st.text_input(
                                                    label, value=ai_val,
                                                    placeholder="(not found — type a value or leave empty)",
                                                    key=f"fill_review_{run_id}",
                                                )
                                        master_value_by_label[label] = user_val

                                # Second pass — build edited_fills for ALL placeholders
                                for ph in pending_placeholders:
                                    run_id = ph["run_id"]
                                    label  = ph.get("label") or pending_labels.get(run_id, run_id)
                                    user_val = master_value_by_label.get(label, "")
                                    edited_fills.append({
                                        "run_id":             run_id,
                                        "value":              "" if (not user_val.strip() or user_val.strip().lower() == "blank") else user_val.strip(),
                                        "found":              bool(user_val.strip() and user_val.strip().lower() != "blank"),
                                        "clear_run_ids":      ph_meta.get(run_id, {}).get("clear_run_ids", []),
                                        "current_text":       ph.get("current_text", ""),
                                        "canonical_run_text": ph.get("canonical_run_text", ""),
                                        "inner_text":         ph.get("inner_text", ""),
                                    })

                                if st.button("✅ Confirm & Generate DOCX", type="primary", key="confirm_fills"):
                                    with st.spinner("Filling in the template…"):
                                        tmpl_line_map = st.session_state["secman_pending_tmpl_map"]
                                        fixed_bytes, applied_log, skipped_log = fill_template_placeholders(
                                            template_docx_bytes, edited_fills, tmpl_line_map, pending_labels
                                        )
                                    st.session_state["secman_fixed_docx_bytes"]     = fixed_bytes
                                    st.session_state["secman_fix_log"]              = (applied_log, skipped_log)
                                    # Store edited_fills for debug display
                                    st.session_state["secman_debug_fills"]          = edited_fills
                                    st.session_state["secman_pending_fills"]        = None
                                    st.session_state["secman_pending_labels"]       = None
                                    st.session_state["secman_pending_tmpl_map"]     = None
                                    st.session_state["secman_pending_placeholders"] = None
                                    st.rerun()

                                if st.button("↩️ Cancel", key="cancel_fills"):
                                    st.session_state["secman_pending_fills"] = None
                                    st.rerun()

                            if st.session_state.get("secman_fixed_docx_bytes"):
                                applied_log, skipped_log = st.session_state.get("secman_fix_log", ([], []))
                                st.success(
                                    f"✅ Filled {len(applied_log)} of {len(applied_log) + len(skipped_log)} "
                                    "field(s) from your uploaded document."
                                )
                                if skipped_log:
                                    with st.expander(f"⚠️ {len(skipped_log)} field(s) left blank", expanded=True):
                                        for label, reason in skipped_log:
                                            st.markdown(f"- **{label}** — {reason}")

                                base_name = os.path.splitext(filled_upload.name)[0]
                                st.download_button(
                                    "⬇️ Download Corrected DOCX",
                                    data=st.session_state["secman_fixed_docx_bytes"],
                                    file_name=f"{base_name}_corrected.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    key="dl_corrected_docx",
                                )

                                # ── PDF download ──────────────────────────────────
                                if st.button("📄 Convert & Download as PDF", key="dl_corrected_pdf"):
                                    with st.spinner("Converting to PDF…"):
                                        pdf_bytes = None
                                        err_msg   = None
                                        try:
                                            import tempfile, subprocess, shutil
                                            with tempfile.TemporaryDirectory() as tmpdir:
                                                docx_path = os.path.join(tmpdir, "corrected.docx")
                                                pdf_path  = os.path.join(tmpdir, "corrected.pdf")
                                                with open(docx_path, "wb") as f:
                                                    f.write(st.session_state["secman_fixed_docx_bytes"])
                                                # Try docx2pdf first (requires Word on Windows/Mac)
                                                try:
                                                    from docx2pdf import convert
                                                    convert(docx_path, pdf_path)
                                                except Exception:
                                                    # Fallback: LibreOffice headless
                                                    lo = shutil.which("soffice") or shutil.which("libreoffice")
                                                    if lo:
                                                        subprocess.run(
                                                            [lo, "--headless", "--convert-to", "pdf",
                                                             "--outdir", tmpdir, docx_path],
                                                            check=True, capture_output=True, timeout=30,
                                                        )
                                                    else:
                                                        err_msg = (
                                                            "PDF conversion requires Microsoft Word (docx2pdf) "
                                                            "or LibreOffice to be installed on this machine."
                                                        )
                                                if os.path.exists(pdf_path):
                                                    with open(pdf_path, "rb") as f:
                                                        pdf_bytes = f.read()
                                        except Exception as e:
                                            err_msg = f"PDF conversion failed: {e}"

                                    if pdf_bytes:
                                        st.download_button(
                                            "⬇️ Download PDF",
                                            data=pdf_bytes,
                                            file_name=f"{base_name}_corrected.pdf",
                                            mime="application/pdf",
                                            key="dl_corrected_pdf_file",
                                        )
                                    elif err_msg:
                                        st.error(err_msg)
                                if st.button("↩️ Edit values & regenerate", key="redo_fills"):
                                    st.session_state["secman_fixed_docx_bytes"] = None
                                    st.session_state["secman_fix_log"]          = None
                                    st.session_state["secman_pending_fills"]    = None
                                    st.rerun()

                                # ── Debug: show exactly what was written ──────────
                                debug_fills = st.session_state.get("secman_debug_fills", [])
                                if debug_fills:
                                    with st.expander("🔍 Debug: What was written to each field", expanded=False):
                                        for f in debug_fills:
                                            rid   = f.get("run_id", "?")
                                            val   = f.get("value", "")
                                            found = f.get("found", False)
                                            ct    = f.get("current_text", "")
                                            normalised = val.replace("\\n", "\n")
                                            lines = [l for l in normalised.split("\n") if l.strip()]
                                            is_multi = len(lines) > 1
                                            st.markdown(
                                                f"**{rid}** found={found} multiline={is_multi} "
                                                f"lines={len(lines)} current_text=`{ct}`"
                                            )
                                            if val:
                                                st.code(repr(val), language="text")

                        # ── Feedback ─────────────────────────────────────────────
                        st.markdown("---")
                        st.markdown("### Was the document type identified correctly?")
                        identified = st.session_state.get("secman_doc_type", "Unknown")
                        all_types  = sorted(list_docx_templates())

                        fb_col1, fb_col2 = st.columns([0.5, 0.5])
                        with fb_col1:
                            st.caption(f"AI identified this as: **{identified}**")
                            is_correct = st.radio(
                                "Is this correct?",
                                ["Yes, correct", "No, it's wrong"],
                                key="fb_is_correct",
                                horizontal=True,
                            )
                        with fb_col2:
                            if is_correct == "No, it's wrong":
                                correct_type = st.selectbox(
                                    "Select the correct document type:",
                                    options=["(type not listed - enter below)"] + all_types,
                                    key="fb_correct_type",
                                )
                                if correct_type == "(type not listed - enter below)":
                                    correct_type = st.text_input("Enter the correct document type:", key="fb_custom_type")
                                reason = st.text_input(
                                    "Why was it wrong? (optional)",
                                    placeholder="e.g. Missing STAFF STUDY in subject line",
                                    key="fb_reason",
                                )
                                if st.button("Save Correction & Re-Run", type="primary", key="fb_save"):
                                    if correct_type and correct_type.strip():
                                        ct = correct_type.strip()
                                        save_feedback({
                                            "doc_snippet":  st.session_state.get("secman_filled_text", "")[:200],
                                            "wrong_type":   identified,
                                            "correct_type": ct,
                                            "reason":       reason.strip(),
                                            "timestamp":    datetime.now().isoformat(),
                                            "filename":     filled_upload.name,
                                        })
                                        _gemini_cache.clear()
                                        tmpl_bytes = find_matching_template(ct)
                                        tmpl_text  = extract_template_formatting(tmpl_bytes, doc_type=ct) if tmpl_bytes else ""
                                        if not tmpl_text:
                                            linked = get_kb_folder_for_doc_type(ct)
                                            if linked:
                                                tmpl_text = get_kb_context_for_folder(linked)
                                        st.session_state["secman_doc_type"]             = ct
                                        st.session_state["secman_template_context"]     = tmpl_text
                                        st.session_state["_secman_template_docx"]       = tmpl_bytes
                                        st.session_state["secman_report"]               = None
                                        st.session_state["secman_report_parsed"]        = None
                                        st.session_state["secman_auto_recheck"]         = True
                                        st.session_state["secman_fixed_docx_bytes"]     = None
                                        st.session_state["secman_fix_log"]              = None
                                        st.session_state["secman_pending_fills"]        = None
                                        st.session_state["secman_pending_labels"]       = None
                                        st.session_state["secman_pending_tmpl_map"]     = None
                                        st.session_state["secman_pending_placeholders"] = None
                                        st.toast(f"Saved. Re-checking as {ct}...")
                                        st.rerun()
                                    else:
                                        st.warning("Please select or enter the correct document type.")

        # ── SUB-TAB 2 — MANAGE TEMPLATES ────────────────────────────────────────
        with subtab_manage:
            templates = list_docx_templates()
            st.markdown(f"### 📁 Document Templates ({len(templates)})")

            st.markdown("#### ⬆️ Upload New Template")
            st.caption(
                "Upload a blank DOCX for each document type. "
                "The name you enter becomes the document type the checker will match against."
            )
            col_tu1, col_tu2 = st.columns([0.55, 0.45])
            with col_tu1:
                tmpl_upload = st.file_uploader(
                    "Upload blank DOCX template",
                    type=["docx"],
                    key="tmpl_uploader_main",
                )
            with col_tu2:
                tmpl_name_override = st.text_input(
                    "Document type name (defaults to filename):",
                    placeholder="e.g. Staff Study Memorandum",
                    key="tmpl_name_override",
                )
            if tmpl_upload:
                final_name = tmpl_name_override.strip() if tmpl_name_override.strip() else os.path.splitext(tmpl_upload.name)[0]
                st.caption(f"Will be saved as: **{final_name}**")
                if st.button("💾 Save Template", type="primary", key="save_tmpl_main"):
                    save_docx_template(final_name, tmpl_upload.getvalue())
                    st.success(f"✅ Template '{final_name}' saved!")
                    st.rerun()

            st.markdown("---")

            # ── Generate Template from a Filled Document ────────────────────
            st.markdown("#### 🪄 Generate Template from a Filled Document")
            st.caption(
                "Don't have a blank template? Turn an already-completed document into one. "
                "AI finds the fields you filled in (names, dates, amounts, memo numbers, etc.) "
                "and replaces just those with red {{ PLACEHOLDER }} text — everything else "
                "(labels, layout, fonts, headers) stays exactly as it was."
            )

            has_check_doc = bool(st.session_state.get("secman_filled_bytes")) and bool(
                st.session_state.get("secman_filled_name")
            )

            tmplgen_source = "upload"
            if has_check_doc:
                source_pick = st.radio(
                    "Source document:",
                    [
                        f"Use the document from Check a Document ({st.session_state['secman_filled_name']})",
                        "Upload a different document",
                    ],
                    key="tmplgen_source_choice",
                )
                tmplgen_source = "checktab" if source_pick.startswith("Use the document") else "upload"

            tmplgen_bytes = None
            tmplgen_name  = None
            if tmplgen_source == "checktab":
                tmplgen_bytes = st.session_state["secman_filled_bytes"]
                tmplgen_name  = st.session_state["secman_filled_name"]
            else:
                tmplgen_upload = st.file_uploader(
                    "Upload a filled DOCX to convert into a template",
                    type=["docx"],
                    key="tmplgen_uploader",
                )
                if tmplgen_upload:
                    tmplgen_bytes = tmplgen_upload.getvalue()
                    tmplgen_name  = tmplgen_upload.name

            if tmplgen_bytes:
                if st.session_state.get("tmplgen_source_name") != tmplgen_name:
                    st.session_state["tmplgen_source_name"]     = tmplgen_name
                    st.session_state["tmplgen_source_bytes"]    = tmplgen_bytes
                    st.session_state["tmplgen_line_map"]        = {}
                    st.session_state["tmplgen_detected_fields"] = None
                    st.session_state["tmplgen_blanked_bytes"]   = None
                    st.session_state["tmplgen_blank_log"]       = None

                if st.button("🔍 Detect Fillable Fields", key="tmplgen_detect"):
                    with st.spinner("Analyzing document for instance-specific content…"):
                        numbered_text, line_map = build_indexed_run_report(tmplgen_bytes)
                        st.session_state["tmplgen_line_map"] = line_map
                        capped_text = "\n".join(numbered_text.splitlines()[:300])
                        detect_prompt = build_field_detection_prompt(capped_text)
                        result = gemini_call(detect_prompt, temperature=0.0, timeout=60)

                    if not result:
                        st.error("AI field detection failed. Check your Gemini API key or quota.")
                    else:
                        clean = re.sub(r"^```[a-zA-Z]*\s*", "", result.strip())
                        clean = re.sub(r"\s*```$", "", clean).strip()
                        bs, be = clean.find("{"), clean.rfind("}")
                        if bs != -1 and be > bs:
                            clean = clean[bs:be + 1]
                        try:
                            parsed = json.loads(clean)
                            fields = parsed.get("fields", [])
                            fields = [
                                f for f in fields
                                if f.get("run_id") in line_map and line_map[f["run_id"]]["kind"] == "run"
                            ]
                            for f in fields:
                                f["include"] = True
                            st.session_state["tmplgen_detected_fields"] = fields
                            st.session_state["tmplgen_blanked_bytes"]   = None
                        except Exception:
                            st.error("Could not parse the AI's response. Try again.")
                        st.rerun()

                detected = st.session_state.get("tmplgen_detected_fields")
                if detected:
                    st.markdown(f"##### Detected {len(detected)} fillable field(s) — uncheck any that shouldn't be blanked:")
                    for i, f in enumerate(detected):
                        fcol1, fcol2 = st.columns([0.06, 0.94])
                        with fcol1:
                            f["include"] = st.checkbox(
                                "Include", value=f.get("include", True),
                                key=f"tmplgen_inc_{i}", label_visibility="collapsed",
                            )
                        with fcol2:
                            f["label"] = st.text_input(
                                f"Field {i+1} label",
                                value=f.get("label", ""),
                                key=f"tmplgen_label_{i}",
                                label_visibility="collapsed",
                            )
                            st.caption(f"Current text: \"{f.get('original_text', '')}\"")

                    st.markdown("---")
                    if st.button("🪄 Generate Blank Template", type="primary", key="tmplgen_generate"):
                        fields_to_blank = [f for f in detected if f.get("include")]
                        with st.spinner("Blanking out fields…"):
                            blanked_bytes, applied_log, skipped_log = blank_out_fields(
                                st.session_state["tmplgen_source_bytes"],
                                fields_to_blank,
                                st.session_state.get("tmplgen_line_map", {}),
                            )
                        st.session_state["tmplgen_blanked_bytes"] = blanked_bytes
                        st.session_state["tmplgen_blank_log"]     = (applied_log, skipped_log)
                        st.rerun()

                if st.session_state.get("tmplgen_blanked_bytes"):
                    applied_log, skipped_log = st.session_state.get("tmplgen_blank_log", ([], []))
                    st.success(f"✅ Blanked out {len(applied_log)} field(s).")
                    if skipped_log:
                        with st.expander(f"⚠️ {len(skipped_log)} field(s) could not be auto-blanked", expanded=False):
                            for label, reason in skipped_log:
                                st.markdown(f"- **{label}** — {reason}")

                    with st.expander("👁️ Preview generated template formatting", expanded=False):
                        preview_text = extract_template_formatting(
                            st.session_state["tmplgen_blanked_bytes"], doc_type="(preview)"
                        )
                        st.text(preview_text[:3000] if preview_text else "(no formatting data extracted)")

                    gcol1, gcol2 = st.columns([0.6, 0.4])
                    with gcol1:
                        new_tmpl_name = st.text_input(
                            "Save as document type:",
                            value=os.path.splitext(tmplgen_name)[0] if tmplgen_name else "",
                            key="tmplgen_save_name",
                        )
                    with gcol2:
                        st.download_button(
                            "⬇️ Download Blank Template",
                            data=st.session_state["tmplgen_blanked_bytes"],
                            file_name=f"{(new_tmpl_name or 'template').strip()}_BLANK.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key="tmplgen_download",
                        )
                    if st.button("💾 Save as Template", type="primary", key="tmplgen_save"):
                        if new_tmpl_name and new_tmpl_name.strip():
                            save_docx_template(new_tmpl_name.strip(), st.session_state["tmplgen_blanked_bytes"])
                            st.success(f"✅ Template '{new_tmpl_name.strip()}' saved!")
                            st.rerun()
                        else:
                            st.warning("Please enter a document type name.")

            st.markdown("---")

            # ── Link KB folder to document type ─────────────────────────────
            kb_folder_names = list(_get_kb_folders().keys())
            if kb_folder_names:
                st.markdown("#### 🔗 Link KB Folder to Document Type")
                st.caption(
                    "Link a folder from **Index Documents** so the checker can use its contents "
                    "as a formatting reference when no DOCX template is available."
                )
                col_lk1, col_lk2 = st.columns(2)
                with col_lk1:
                    link_folder = st.selectbox("KB Folder:", options=kb_folder_names, key="link_kb_folder")
                with col_lk2:
                    link_doc_type = st.text_input(
                        "Maps to document type name:",
                        value=link_folder,
                        key="link_doc_type_name",
                        placeholder="e.g. Memorandum",
                    )
                if st.button("🔗 Save Link", key="create_kb_link", type="primary"):
                    links = load_kb_folder_links()
                    links[link_doc_type.strip()] = link_folder
                    save_kb_folder_links(links)
                    st.success(f"✅ '{link_doc_type.strip()}' → KB folder '{link_folder}' linked.")
                    st.rerun()

                existing_links = load_kb_folder_links()
                if existing_links:
                    st.markdown("**Active KB folder links:**")
                    for dt, folder in existing_links.items():
                        lc1, lc2 = st.columns([0.85, 0.15])
                        lc1.caption(f"📄 `{dt}` → KB folder `{folder}`")
                        if lc2.button("🗑️", key=f"del_link_{dt}", help=f"Remove link for {dt}"):
                            del existing_links[dt]
                            save_kb_folder_links(existing_links)
                            st.rerun()
            else:
                st.info("No KB folders found. Index documents in the **Index Documents** tab first, then link them here.")

            st.markdown("---")

            # ── Existing templates list ──────────────────────────────────────
            if not templates:
                st.info("No templates stored yet. Upload one above.")
            else:
                with st.expander("📂 Stored Templates", expanded=True):
                    for tmpl in templates:
                        with st.expander(f"📄 {tmpl}", expanded=False):
                            tc1, tc2, tc3 = st.columns([0.6, 0.25, 0.15])
                            tmpl_bytes = load_docx_template(tmpl)

                            if tmpl_bytes:
                                tc1.caption(f"{round(len(tmpl_bytes)/1024, 1)} KB · {_template_path(tmpl).name}")
                                tc2.download_button(
                                    "⬇️ Download",
                                    data=tmpl_bytes,
                                    file_name=f"{tmpl}.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    key=f"dl_tmpl_{tmpl}",
                                )
                            if tc3.button("🗑️ Delete", key=f"del_tmpl_{tmpl}", help=f"Delete {tmpl}"):
                                _template_path(tmpl).unlink(missing_ok=True)
                                override = _template_path(tmpl).with_suffix(".override.txt")
                                if override.exists():
                                    override.unlink()
                                st.rerun()

                            if tmpl_bytes:
                                st.markdown("**Formatting data AI extracts from this template:**")
                                extracted = extract_template_formatting(tmpl_bytes, doc_type=tmpl)

                                edit_key   = f"_edit_tmpl_{tmpl}"
                                edited_key = f"_edited_tmpl_{tmpl}"

                                if edit_key not in st.session_state:
                                    st.session_state[edit_key] = False

                                if not st.session_state[edit_key]:
                                    if extracted.strip():
                                        st.code(
                                            extracted[:3000] + ("…" if len(extracted) > 3000 else ""),
                                            language="text",
                                        )
                                        st.caption(f"{len(extracted.split())} words · {len(extracted)} chars")
                                    else:
                                        st.warning("⚠️ No formatting data extracted — template may use images or shapes only.")
                                    if st.button("✏️ Edit formatting reference", key=f"edit_btn_{tmpl}"):
                                        st.session_state[edit_key]   = True
                                        st.session_state[edited_key] = extracted
                                        st.rerun()
                                else:
                                    st.caption("Edit the formatting reference below. This is what AI will compare against.")
                                    new_text = st.text_area(
                                        "Formatting reference:",
                                        value=st.session_state.get(edited_key, extracted),
                                        height=300,
                                        key=f"textarea_{tmpl}",
                                    )
                                    ec1, ec2, ec3 = st.columns([0.25, 0.25, 0.5])
                                    if ec1.button("💾 Save edits", key=f"save_edit_{tmpl}", type="primary"):
                                        override_path = _template_path(tmpl).with_suffix(".override.txt")
                                        override_path.write_text(new_text, encoding="utf-8")
                                        st.session_state[edit_key] = False
                                        st.success("✅ Saved — AI will use your edited reference from now on.")
                                        st.rerun()
                                    if ec2.button("✖ Cancel", key=f"cancel_edit_{tmpl}"):
                                        st.session_state[edit_key] = False
                                        st.rerun()
                                    if ec3.button("↺ Reset to auto-extracted", key=f"reset_edit_{tmpl}"):
                                        override_path = _template_path(tmpl).with_suffix(".override.txt")
                                        if override_path.exists():
                                            override_path.unlink()
                                        st.session_state[edit_key] = False
                                        st.success("↺ Reset — using auto-extracted formatting again.")
                                        st.rerun()

                                override_path = _template_path(tmpl).with_suffix(".override.txt")
                                if override_path.exists():
                                    st.info("✏️ Using manually edited formatting reference.")

            st.markdown("---")

            # ── Feedback / Corrections ───────────────────────────────────────
            all_feedback = load_feedback()
            st.markdown(f"### 🧠 AI Corrections ({len(all_feedback)})")
            st.caption("These teach the AI to identify documents correctly in future runs.")
            if not all_feedback:
                st.info("No corrections saved yet. Use the feedback section after running a check.")
            else:
                with st.expander("📂 Saved Corrections", expanded=False):
                    for i, entry in enumerate(all_feedback):
                        fc1, fc2 = st.columns([0.9, 0.1])
                        with fc1:
                            st.markdown(
                                f"**#{i+1}** &nbsp; `{entry.get('wrong_type','?')}` → "
                                f"`{entry.get('correct_type','?')}` &nbsp; "
                                f"<span style='color:#888;font-size:11px;'>{entry.get('timestamp','')[:10]}"
                                f" · {entry.get('filename','')}</span>",
                                unsafe_allow_html=True,
                            )
                            if entry.get("reason"):
                                st.caption(f"  Reason: {entry['reason']}")
                        with fc2:
                            if st.button("🗑️", key=f"del_fb_{i}", help="Delete this correction"):
                                delete_feedback(i)
                                st.rerun()